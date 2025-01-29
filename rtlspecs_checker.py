import os
import base64
import olefile
import io
import openai
import csv
import docx
import textwrap
from pyverilog.vparser.parser import parse
from pyverilog.vparser.ast import (
    Decl, Reg, Localparam, Parameter, IntConst, Identifier,
    UnaryOperator, Operator, Partselect, Pointer, NonblockingSubstitution,
    BlockingSubstitution, IfStatement, CaseStatement, Block, Rvalue,
    Input, Output, Inout, Ioport, Port,
    InstanceList, Instance, Always, ForStatement, Assign
)
from docx import Document
from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image
from openai import OpenAI
from collections import defaultdict, deque

# Load your OpenAI API key from the environment variable
openai.api_key = os.getenv('OPENAI_API_KEY')

# Set your OpenAI API key
client = openai

# Set to True to enable debug output
DEBUG = False


class ModuleNode:
    def __init__(self, module_name, instance_name=None):
        self.module_name = module_name
        self.instance_name = instance_name
        self.children = []  # List of ModuleNode instances

    def add_child(self, child_node):
        self.children.append(child_node)

    def print_hierarchy(self, level=0):
        indent = "    " * level
        if self.instance_name:
            print(f"{indent}- Instance: {self.instance_name}, Module: {self.module_name}")
        else:
            print(f"Top-Level Module: {self.module_name}")
        for child in self.children:
            child.print_hierarchy(level + 1)


def get_file_list(directory, extension=".v"):
    """
    Returns a list of Verilog files in the specified directory.

    :param directory: Directory to search for Verilog files.
    :param extension: File extension to filter by (default is ".v").
    :return: List of Verilog file paths.
    """
    return [os.path.join(directory, f) for f in os.listdir(directory) if f.endswith(extension)]


def generate_asts(file_list):
    """
    Parses Verilog files and returns their ASTs.

    :param file_list: List of Verilog file paths.
    :return: List of ASTs corresponding to the Verilog files.
    """
    asts = []
    for filename in file_list:
        try:
            ast, _ = parse([filename])
            asts.append(ast)
        except Exception as e:
            print(f"[ERROR] Error parsing {filename}: {e}")
    return asts


def print_ast_contents(asts):
    """
    Prints the AST contents for each module.

    :param asts: List of ASTs corresponding to the Verilog files.
    """
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                print(f"\nAST for module: {module_name}")
                print_ast_node(definition, "")


def print_ast_node(node, indent):
    """
    Recursively prints the AST node and its children with enhanced readability.

    :param node: The AST node to print.
    :param indent: The indentation level for pretty printing.
    """
    if isinstance(node, Node):
        node_info = f"{indent}{node.__class__.__name__}"

        # Add more detailed information for specific types of nodes
        if isinstance(node, (Parameter, Localparam)):
            node_info += f" (name={node.name}, value={get_rvalue(node.value)})"
        elif isinstance(node, Reg):
            node_info += f" (name={node.name}, width={get_rvalue(node.width)})"
        elif isinstance(node, Assign):
            node_info += f" (LHS={get_rvalue(node.left)}, RHS={get_rvalue(node.right)})"
        elif isinstance(node, Always):
            sens_list = get_rvalue(node.sens_list) if node.sens_list else 'None'
            node_info += f" (sensitivity={sens_list})"
        elif isinstance(node, IfStatement):
            condition_str = get_rvalue(node.cond)
            node_info += f" (condition={condition_str})"
        elif isinstance(node, CaseStatement):
            comp_expr = get_rvalue(node.comp)
            node_info += f" (expression={comp_expr})"
        else:
            if hasattr(node, 'name'):
                node_info += f" (name={node.name})"
            if hasattr(node, 'value'):
                node_info += f" (value={get_rvalue(node.value)})"

        print(node_info)

        # Recursively print each child node with increased indentation
        for child in node.children():
            print_ast_node(child, indent + "    ")
    else:
        # Print non-Node types directly (e.g., strings, integers)
        print(f"{indent}{node}")


# def get_rvalue(value):
#     if isinstance(value, IntConst):
#         return value.value
#     elif isinstance(value, Identifier):
#         return value.name
#     elif isinstance(value, UnaryOperator):
#         # For unary operators like ~, !
#         return f"{value.op}{get_rvalue(value.right)}"  # or something similar
#     elif isinstance(value, Operator):
#         # Usually a binary operator, e.g. value.left, value.op, value.right
#         left_str = get_rvalue(value.left)
#         right_str = get_rvalue(value.right)
#         op_str = get_operator_symbol_by_op(value.op)
#         return f"({left_str} {op_str} {right_str})"
#     elif isinstance(value, Partselect):
#         var = get_rvalue(value.var)
#         msb = get_rvalue(value.msb)
#         lsb = get_rvalue(value.lsb)
#         return f"{var}[{msb}:{lsb}]"
#     elif isinstance(value, Pointer):
#         var = get_rvalue(value.var)
#         ptr = get_rvalue(value.ptr)
#         return f"{var}[{ptr}]"
#     else:
#         return str(value)

def get_rvalue(value):
    if isinstance(value, IntConst):
        return value.value
    elif isinstance(value, Identifier):
        return value.name
    elif isinstance(value, UnaryOperator):
        return f"{value.op}{get_rvalue(value.right)}"
    elif isinstance(value, Operator):
        left_str = get_rvalue(value.left)
        right_str = get_rvalue(value.right)
        op_str = get_operator_symbol_by_op(value.op)
        return f"({left_str} {op_str} {right_str})"
    elif isinstance(value, Partselect):
        var = get_rvalue(value.var)
        msb = get_rvalue(value.msb)
        lsb = get_rvalue(value.lsb)
        return f"{var}[{msb}:{lsb}]"
    elif isinstance(value, Pointer):
        var = get_rvalue(value.var)
        ptr = get_rvalue(value.ptr)
        return f"{var}[{ptr}]"

    elif isinstance(value, Rvalue):
        # Rvalue is basically a simple wrapper with a single child (the actual expression).
        # We can just recurse on its only child:
        children = value.children()
        if len(children) == 1:
            return get_rvalue(children[0])
        else:
            # If there's more than one child (rare for a parameter), handle them as needed.
            # For most Verilog parameters, this won't happen, so we do:
            return ''.join(get_rvalue(ch) for ch in children)

    else:
        # Fallback: Convert unknown AST nodes to string
        return str(value)


def get_operator_symbol_by_op(op_str):
    operator_map = {
        '+': '+',
        '-': '-',
        '*': '*',
        '/': '/',
        '%': '%',
        '==': '==',
        '!=': '!=',
        '>': '>',
        '>=': '>=',
        '<': '<',
        '<=': '<=',
        '&&': '&&',
        '||': '||',
        '&': '&',
        '|': '|',
        '^': '^',
        '^~': '^~',
        '<<': '<<',
        '>>': '>>',
    }
    return operator_map.get(op_str, op_str)


def extract_parameters_from_asts(asts):
    parameters_dict = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                parameters_dict[module_name] = {}
                for item in definition.items:
                    if isinstance(item, Decl):
                        for decl in item.list:
                            if isinstance(decl, Parameter) and not isinstance(decl, Localparam):
                                parameters_dict[module_name][decl.name] = get_rvalue(decl.value)
                    elif isinstance(item, Parameter) and not isinstance(item, Localparam):
                        parameters_dict[module_name][item.name] = get_rvalue(item.value)
    return parameters_dict


def print_parameters(parameters_dict):
    """
    Prints the parameters for each module in CSV format.

    :param parameters_dict: Dictionary with module names as keys and their parameters as values.
    """
    for module_name, params in parameters_dict.items():
        print(f"\nRaw Parameters for module {module_name}:\n")
        if params:
            # Print CSV header
            print("Name,Value")
            for param_name, param_value in params.items():
                print(f"{param_name},{param_value}")
        else:
            print("No parameters found.")


def extract_ioports_from_asts(asts):
    """
    Extracts IO ports' information from the ASTs.
    Returns a dictionary with module names as keys and a list of dictionaries with keys 'Name', 'Width', 'Direction'.

    :param asts: List of ASTs corresponding to the Verilog files.
    :return: Dictionary with module names as keys and their IO ports' information as values.
    """
    ioports_dict = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                ioports_list = []
                for item in definition.items:
                    if isinstance(item, Decl):
                        for decl in item.list:
                            if isinstance(decl, (Input, Output, Inout)):
                                port_direction = decl.__class__.__name__.lower()
                                port_name = decl.name
                                port_width = extract_port_width(decl)
                                width = port_width if port_width else '1'  # If no width specified, assume width is '1'
                                # Append a dictionary
                                ioports_list.append({
                                    'Name': port_name,
                                    'Width': width,
                                    'Direction': port_direction
                                })
                ioports_dict[module_name] = ioports_list
    return ioports_dict


def extract_port_width(decl):
    """
    Extracts the width of a port as a string.

    :param decl: The declaration node of the port.
    :return: A string representing the width of the port.
    """
    if hasattr(decl, 'width') and decl.width is not None:
        msb = get_rvalue(decl.width.msb)
        lsb = get_rvalue(decl.width.lsb)
        # Try to compute the width numerically
        try:
            msb_int = int(msb)
            lsb_int = int(lsb)
            width = abs(msb_int - lsb_int) + 1
            return str(width)
        except ValueError:
            # msb or lsb is not an integer, return as is
            return f"[{msb}:{lsb}]"
    else:
        return "1"  # Default width for 1-bit signals


def print_ioports(ioports_dict):
    """
    Prints the IO ports' information for each module in the same format as extract_ioports_from_specs.

    :param ioports_dict: Dictionary with module names as keys and their IO ports' information as values.
    """
    for module_name, ioports in ioports_dict.items():
        print(f"\nRaw IO Ports for module {module_name}:\n")
        if ioports:
            # Print CSV header
            print("Name,Width,Direction")
            for port in ioports:
                print(f"{port['Name']},{port['Width']},{port['Direction']}")
        else:
            print("No IO ports found.")

def extract_module_hierarchy(asts):
    # Build a mapping from module names to their AST definitions
    module_definitions = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                module_definitions[module_name] = definition

    # Build a mapping from modules to the modules they instantiate
    instantiated_modules = {}
    module_instantiations = {}
    for module_name, definition in module_definitions.items():
        instantiated = []
        instantiations = []
        for item in definition.items:
            if isinstance(item, InstanceList):
                for instance in item.instances:
                    if isinstance(instance, Instance):
                        child_module_name = instance.module
                        instance_name = instance.name
                        instantiated.append(child_module_name)
                        instantiations.append((child_module_name, instance_name))
        instantiated_modules[module_name] = instantiated
        module_instantiations[module_name] = instantiations

    # Identify top-level modules (modules not instantiated by any other module)
    all_modules = set(module_definitions.keys())
    instantiated_modules_set = set()
    for modules in instantiated_modules.values():
        instantiated_modules_set.update(modules)
    top_level_modules = all_modules - instantiated_modules_set

    # Recursively build the module hierarchy starting from top-level modules
    hierarchy_trees = []
    for top_module_name in top_level_modules:
        root_node = ModuleNode(module_name=top_module_name)
        build_hierarchy_recursive(root_node, module_instantiations, module_definitions)
        hierarchy_trees.append(root_node)

    return hierarchy_trees

def build_hierarchy_recursive(current_node, module_instantiations, module_definitions, visited=None):
    if visited is None:
        visited = set()
    module_name = current_node.module_name
    if module_name in visited:
        # Prevent infinite recursion in case of cyclic dependencies
        return
    visited.add(module_name)
    instantiations = module_instantiations.get(module_name, [])
    for child_module_name, instance_name in instantiations:
        child_node = ModuleNode(module_name=child_module_name, instance_name=instance_name)
        current_node.add_child(child_node)
        # Recursively build the hierarchy for the child module
        build_hierarchy_recursive(child_node, module_instantiations, module_definitions, visited.copy())

def print_module_hierarchy(hierarchy_trees):
    for tree in hierarchy_trees:
        tree.print_hierarchy()
        print("")  # Add a blank line between top-level modules


def extract_text_by_chapters(docx_path):
    """
    Extracts text from a Word document and chunks it based on chapter titles,
    including tables, images, and equations, while maintaining the correct order of elements.
    """
    doc = Document(docx_path)
    chapters = []
    current_chapter = {'title': '', 'content': ''}

    image_counter = 1
    equation_counter = 1  # Initialize equation_counter here

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            para = block
            if para.style.name == 'Heading 1':
                if current_chapter['title'] or current_chapter['content']:
                    chapters.append(current_chapter)
                    current_chapter = {'title': '', 'content': ''}  # Start new chapter
                current_chapter['title'] = para.text.strip()
            else:
                if has_equation(para):
                    # Process paragraph with equations
                    equation_counter = process_paragraph_with_equations(para, current_chapter, equation_counter)
                else:
                    # Process runs in the paragraph
                    for run in para.runs:
                        if has_image(run):
                            image_filename = f"image_{image_counter}.png"
                            extract_image(run, image_filename)
                            image_counter += 1
                            current_chapter['content'] += f"[Image: {image_filename}]\n"
                        else:
                            # Add run text
                            current_chapter['content'] += run.text
                    current_chapter['content'] += '\n'
        elif isinstance(block, Table):
            table = block
            # Extract text from the table
            table_text = extract_text_from_table(table)
            # Add table text to the current chapter's content
            current_chapter['content'] += table_text + '\n'
        else:
            # Handle other types of blocks if needed
            pass
    # Add the last chapter
    if current_chapter['title'] or current_chapter['content']:
        chapters.append(current_chapter)
    return chapters


def has_equation(para):
    """
    Checks if the paragraph contains an equation, either as oMath elements or oMathPara elements.
    """
    # Define namespaces
    MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

    # Define fully qualified element names
    oMath_tag = f'{{{MATH_NS}}}oMath'
    oMathPara_tag = f'{{{MATH_NS}}}oMathPara'

    oMath_elements = para._element.findall('.//' + oMath_tag)
    oMathPara_elements = para._element.findall('.//' + oMathPara_tag)

    return bool(oMath_elements or oMathPara_elements)


def process_paragraph_with_equations(para, current_chapter, equation_counter):
    """
    Processes a paragraph containing equations, extracting text and equations
    in order and updating the current chapter's content.
    """
    MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    oMath_tag = f'{{{MATH_NS}}}oMath'
    oMathPara_tag = f'{{{MATH_NS}}}oMathPara'

    content_elements = []

    for child in para._element.iterchildren():
        if child.tag in (oMath_tag, oMathPara_tag):
            # It's an equation
            equation_filename = f"equation_{equation_counter}.xml"
            with open(equation_filename, 'w', encoding='utf-8') as f:
                xml_content = etree.tostring(child, pretty_print=True, encoding='unicode')
                f.write(xml_content)
            content_elements.append(f"[Equation: {equation_filename}]")
            equation_counter += 1
        else:
            # It's text
            text = ''.join(child.itertext())
            content_elements.append(text)

    # Combine the content elements into the current chapter's content
    current_chapter['content'] += ''.join(content_elements) + '\n'

    return equation_counter


def has_image(run):
    """
    Checks if the run contains an image.
    """
    blip_elements = run._element.findall('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
    if blip_elements:
        return True
    # Check for embedded images in v:imagedata (used for OLE objects)
    imagedata_elements = run._element.findall('.//v:imagedata', {'v': 'urn:schemas-microsoft-com:vml'})
    if imagedata_elements:
        return True
    return False


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within parent,
    in document order. Each returned value is an instance of either Table or Paragraph.

    :param parent: The parent object (Document or _Cell).
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type.")

    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)
        else:
            # Handle other types of elements if needed
            pass


def extract_text_from_table(table):
    """
    Extracts text from a table and formats it in Markdown table syntax.

    :param table: The Table object.
    :return: A string representation of the table's content in Markdown format.
    """
    md_table = []

    # Extract the headers (assuming the first row contains headers)
    headers = []
    first_row = table.rows[0]
    for cell in first_row.cells:
        cell_text = ''
        for paragraph in cell.paragraphs:
            cell_text += paragraph.text.strip() + ' '
        headers.append(cell_text.strip())

    # Build the header row
    header_row = '| ' + ' | '.join(headers) + ' |'
    md_table.append(header_row)

    # Build the separator row
    separator_row = '| ' + ' | '.join(['---'] * len(headers)) + ' |'
    md_table.append(separator_row)

    # Process the remaining rows
    for row in table.rows[1:]:
        row_cells = []
        for cell in row.cells:
            cell_text = ''
            for paragraph in cell.paragraphs:
                cell_text += paragraph.text.strip() + ' '
            row_cells.append(cell_text.strip())
        data_row = '| ' + ' | '.join(row_cells) + ' |'
        md_table.append(data_row)

    # Join all rows into a single string
    table_text = '\n'.join(md_table)
    return table_text


def extract_image(run, image_filename):
    """
    Extracts an image from a run and saves it to a file.

    :param run: The run containing the image.
    :param image_filename: The filename to save the image as.
    """
    # Define namespaces
    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'v': 'urn:schemas-microsoft-com:vml',
    }

    # Find all 'blip' elements within the run
    blip_elements = run._element.findall('.//a:blip', nsmap)
    for blip in blip_elements:
        rEmbed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rEmbed:
            image_part = run.part.related_parts[rEmbed]
            with open(image_filename, 'wb') as f:
                f.write(image_part.blob)
            return  # Image extracted

    # Check for 'imagedata' elements (used in OLE objects)
    imagedata_elements = run._element.findall('.//v:imagedata', nsmap)
    for imagedata in imagedata_elements:
        rId = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        if rId:
            image_part = run.part.related_parts[rId]
            with open(image_filename, 'wb') as f:
                f.write(image_part.blob)
            return  # Image extracted


def extract_parameters_from_specs(chapter, module_name):
    """
    Extracts RTL default parameters from the specifications using GPT-4o-mini for a specified module name.
    Directly prints the GPT output for the module.

    :param chapter: A dictionary with 'title' and 'content' of the 'RTL Parameters' chapter.
    :param module_name: The module name to extract parameters for.
    """
    title = chapter['title']
    content = chapter['content']

    # Prepare the prompt
    prompt = (
        f"Please extract the default RTL parameters' value for the module {module_name} from the following 'RTL Parameters' chapter. "
        f"Exclude any submodule RTL default parameters and focus only on module {module_name}. Provide the details in CSV format with the headers 'Name' and 'Value'. "
        f"Do not include any additional text or explanations. Only provide the CSV content. For example:\n\n"
        f"Name,Value\nPARAM1,VALUE1\nPARAM2,VALUE2\n\n"
        f"Now, extract the RTL default parameters:\n\n"
        f"Title: {title}\n\nContent:\n{content}\n"
    )

    # Call GPT-4o-mini to extract parameters
    try:
        completion = client.chat.completions.create(
            model='gpt-4o-mini',
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an assistant that extracts default RTL parameters' value from specifications. "
                        "When provided with a module name and specifications content, you should output only the CSV data containing "
                        "the parameters' name and their default RTL parameters' value for that module. Do not include any explanations or additional text."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0,      # For deterministic output
            max_tokens=1000,    # Increase as needed
            stop = ["\n\n"]     # Stop after two newlines (adjust as needed)
        )

        # Get the assistant's reply
        assistant_reply = completion.choices[0].message.content

        # Print the raw assistant reply for debugging
        if DEBUG:
            print(f"\nRaw Assistant Reply for module {module_name}:\n")
            print(assistant_reply)

        # Parse the parameters into a dictionary
        csv_content = assistant_reply.strip()
        reader = csv.DictReader(io.StringIO(csv_content))
        parameters = {row['Name']: row['Value'] for row in reader}

        return parameters

    except Exception as e:
        print(f"[ERROR] Error extracting parameters for module {module_name}: {e}")
        return {}


def normalize_string(s):
    """
    Normalizes a string by replacing curly quotes with straight quotes and removing extra whitespace.
    """
    if s is None:
        return ''
    # Replace curly single quotes with straight quotes
    s = s.replace('‘', "'").replace('’', "'")
    # Replace curly double quotes with straight quotes (if applicable)
    s = s.replace('“', '"').replace('”', '"')
    # Replace any non-breaking spaces with regular spaces (optional)
    s = s.replace('\u00A0', ' ')
    # Strip leading and trailing whitespace
    s = s.strip()
    return s


def compare_parameters(rtl_params_dict, spec_params_dict):
    """
    Compares parameters from RTL and specifications.
    Prints matched and mismatched parameters and summarizes the results.

    :param rtl_params_dict: Dictionary of parameters from RTL with module names as keys.
    :param spec_params_dict: Dictionary of parameters from specifications with module names as keys.
    """
    total_matched = 0
    total_mismatched = 0

    # Define the format string for aligned output
    format_str = "[{:<5}] Status: {:<12} Module: {:<15} Parameter: {:<35} Specs: {:<25} RTL: {}"

    # Get the union of module names from both RTL and specifications
    all_module_names = set(spec_params_dict.keys()) | set(rtl_params_dict.keys())

    for module_name in all_module_names:
        rtl_params = rtl_params_dict.get(module_name, {})
        spec_params = spec_params_dict.get(module_name, {})

        # Get the union of parameter names from both RTL and specifications for this module
        all_param_names = set(spec_params.keys()) | set(rtl_params.keys())

        for param_name in all_param_names:
            rtl_value = normalize_string(rtl_params.get(param_name, 'NA'))
            spec_value = normalize_string(spec_params.get(param_name, 'NA'))

            # Skip parameters missing in both
            if rtl_value == 'N/A' and spec_value == 'N/A':
                continue

            if rtl_value == spec_value:
                print(format_str.format("INFO", "MATCHED", module_name, param_name, spec_value, rtl_value))
                total_matched += 1
            else:
                print(format_str.format("ERROR", "MISMATCHED", module_name, param_name, spec_value, rtl_value))
                total_mismatched += 1

    # Print summary
    return ("RTL Default Parameters Check", f"MATCHED: {total_matched}, MISMATCHED: {total_mismatched}")


def extract_ioports_from_specs(chapter, module_name):
    """
    Extracts IO ports' information from the specifications using GPT for a specified module name.
    Returns a list of dictionaries with keys 'Name', 'Width', 'Direction'.

    :param chapter: A dictionary with 'title' and 'content' of the 'RTL Parameters' chapter.
    :param module_name: The module name to extract parameters for.
    """
    title = chapter['title']
    content = chapter['content']

    # Prepare the prompt
    prompt = (
        f"Please extract the IO ports information for the module {module_name} from the following 'IO Ports' chapter. "
        f"Exclude any submodule IO ports' information and focus only on module {module_name}. Provide the details in CSV format with the headers 'Name','Width','Direction'. "
        f"Do not include any additional text or explanations. Only provide the CSV content. For example:\n\n"
        f"Name,Width,Direction\nPORT1,32,input\nPORT2,1,output\n\n"
        f"Now, extract the IO ports information:\n\n"
        f"Title: {title}\n\nContent:\n{content}\n"
    )

    # Call GPT to extract IO ports' information
    try:
        completion = client.chat.completions.create(
            model='gpt-4o-mini',
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an assistant that extracts IO ports' information from specifications. "
                        "When provided with a module name and specifications content, you should output only the CSV data containing "
                        "the port's name, width, and direction for that module. Do not include any explanations or additional text."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=1000,
            stop=["\n\n"]
        )

        # Get the assistant's reply
        assistant_reply = completion.choices[0].message.content

        # Print the raw assistant reply for debugging
        if DEBUG:
            print(f"\nRaw Assistant Reply for module {module_name}:\n")
            print(assistant_reply)

        # Parse the IO ports' information into a list of dictionaries
        csv_content = assistant_reply.strip()
        reader = csv.DictReader(io.StringIO(csv_content))
        ioports = [row for row in reader]

        return ioports

    except Exception as e:
        print(f"[ERROR] Error extracting IO ports' information for module {module_name}: {e}")
        return []


def compare_ioports(rtl_ioports_dict, spec_ioports_dict):
    """
    Compares IO ports' information from RTL and specifications.
    Prints matched and mismatched IO ports' information and summarizes the results.

    :param rtl_ioports_dict: Dictionary with module names as keys and their IO ports' information as values.
    :param spec_ioports_dict: Dictionary with module names as keys and their IO ports' information as values.
    """
    total_matched = 0
    total_mismatched = 0

    # Define the format string for aligned output
    format_str = "[{:<5}] Status: {:<12} Module: {:<15} Port: {:<20} Specs: Width: {:>2} Direction: {:<6}  RTL: Width: {:>2} Direction: {:<6}"

    # Get the union of module names from both RTL and specifications
    all_module_names = set(spec_ioports_dict.keys()) | set(rtl_ioports_dict.keys())

    for module_name in sorted(all_module_names):
        rtl_ioports = rtl_ioports_dict.get(module_name, [])
        spec_ioports = spec_ioports_dict.get(module_name, [])

        # Create dictionaries keyed by port name for quick lookup
        rtl_ports_dict = {
            normalize_string(port['Name']): {
                'Width': normalize_string(port['Width']),
                'Direction': normalize_string(port['Direction'])
            } for port in rtl_ioports
        }
        spec_ports_dict = {
            normalize_string(port['Name']): {
                'Width': normalize_string(port['Width']),
                'Direction': normalize_string(port['Direction'])
            } for port in spec_ioports
        }

        all_port_names = set(rtl_ports_dict.keys()) | set(spec_ports_dict.keys())

        for port_name in sorted(all_port_names):
            rtl_port = rtl_ports_dict.get(port_name)
            spec_port = spec_ports_dict.get(port_name)

            if rtl_port and spec_port:
                # Port exists in both RTL and Specs
                if rtl_port == spec_port:
                    # Exact match
                    print(
                        format_str.format(
                            "INFO",
                            "MATCHED",
                            module_name,
                            port_name,
                            spec_port['Width'],
                            spec_port['Direction'],
                            rtl_port['Width'],
                            rtl_port['Direction']
                        )
                    )
                    total_matched += 1
                else:
                    # Width or Direction mismatch
                    print(
                        format_str.format(
                            "ERROR",
                            "MISMATCHED",
                            module_name,
                            port_name,
                            spec_port['Width'],
                            spec_port['Direction'],
                            rtl_port['Width'],
                            rtl_port['Direction']
                        )
                    )
                    total_mismatched += 1
            elif rtl_port and not spec_port:
                # Port only in RTL
                print(
                    format_str.format(
                        "ERROR",
                        "MISMATCHED",
                        module_name,
                        port_name,
                        "N/A",
                        "N/A",
                        rtl_port['Width'],
                        rtl_port['Direction']
                    )
                )
                total_mismatched += 1
            elif spec_port and not rtl_port:
                # Port only in Specs
                print(
                    format_str.format(
                        "ERROR",
                        "MISMATCHED",
                        module_name,
                        port_name,
                        spec_port['Width'],
                        spec_port['Direction'],
                        "N/A",
                        "N/A",
                    )
                )
                total_mismatched += 1

    # Print summary
    return ("IO Ports Check", f"MATCHED: {total_matched}, MISMATCHED: {total_mismatched}")


def get_hierarchy_text(node, level=0):
    """
    Generates a text representation of the module hierarchy starting from the given node.

    :param node: ModuleNode instance representing the root of the hierarchy.
    :param level: Current indentation level.
    :return: A string representing the hierarchy.
    """
    indent = "    " * level
    if node.instance_name:
        line = f"{indent}- Instance: {node.instance_name}, Module: {node.module_name}\n"
    else:
        line = f"{indent}Top-Level Module: {node.module_name}\n"
    for child in node.children:
        line += get_hierarchy_text(child, level + 1)
    return line


def compare_module_hierarchy(hierarchy_trees, specs_module_hierarchy_text):
    """
    Compares the module hierarchy extracted from the RTL with the module hierarchy from the specifications.
    Uses GPT to perform the comparison and outputs the result.

    :param hierarchy_trees: List of ModuleNode instances representing the RTL module hierarchy.
    :param specs_module_hierarchy_text: String containing the module hierarchy from the specifications.
    """
    # Convert the hierarchy_trees to text
    rtl_module_hierarchy_text = ''
    for tree in hierarchy_trees:
        rtl_module_hierarchy_text += get_hierarchy_text(tree) + '\n'

    # -----------------------------------------------------------------
    # 1) Print out the hierarchy from RTL and from Specs before calling GPT
    # -----------------------------------------------------------------
    print("RTL Module Hierarchy")
    print("====================")
    print(rtl_module_hierarchy_text.strip())

    print("")
    print("Specs Module Hierarchy")
    print("======================")
    print(specs_module_hierarchy_text.strip())
    print("")

    # Prepare the prompt for GPT
    prompt = (
        "You are an assistant that compares module hierarchies from RTL code and specifications. "
        "You will be provided with two module hierarchies: one extracted from RTL code and one from specifications. "
        "Compare them based on module names, instance names, hierarchy structure and any any textual details, including the number of key instances. "
        "Pay close attention to numerical details and report any discrepancies. "
        "If they match in all aspects, output:\n\n"
        "[INFO ] MATCHED\n\n"
        "If they do not match, output:\n\n"
        "[ERROR] MISMATCHED. <Provide the precise reason for the mismatch.>\n\n"
        "Please clearly output the discrepancies show in rtl and specifications.\n\n"
        "Only output the specified message, without any additional text.\n\n"
        "Here are the module hierarchies:\n\n"
        "RTL Module Hierarchy:\n"
        "{}\n"
        "Specifications Module Hierarchy:\n"
        "{}"
    ).format(rtl_module_hierarchy_text.strip(), specs_module_hierarchy_text.strip())

    if DEBUG:
        print("GPT Prompt:\n", prompt)

    # Call GPT to perform the comparison
    try:
        completion = client.chat.completions.create(
            model='gpt-4o',
            messages=[
                {"role": "system",
                 "content": (
                     "You are an assistant that compares module hierarchies from RTL code and specifications. "
                     "You only output the comparison result precisely as specified."
                 )},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=500,
            n=1,
            stop=None,
        )

        # Get the assistant's reply
        assistant_reply = completion.choices[0].message.content.strip()

        # Print the assistant's reply
        print(assistant_reply)

        # Strip off [INFO] and [ERROR] tags from the response
        if "[INFO ]" in assistant_reply:
            result = assistant_reply.replace("[INFO ]", "").strip()
            return ("Module Hierarchy Check", result)
        elif "[ERROR]" in assistant_reply:
            result = assistant_reply.replace("[ERROR]", "").strip()
            return ("Module Hierarchy Check", result)
        else:
            return ("Module Hierarchy Check", assistant_reply)

    except Exception as e:
        print(f"[ERROR] Error comparing Module Hierarchy from specifications and RTL: {e}")


def is_clock_signal(signal_name):
    """
    Determines if a signal name corresponds to a clock signal based on common naming conventions.
    """
    clock_keywords = ['clk', 'clock']
    signal_name_lower = signal_name.lower()
    return any(keyword in signal_name_lower for keyword in clock_keywords)


class ModuleClockDomainInfo:
    def __init__(self, module_name):
        self.module_name = module_name
        self.clocks_in_module = []    # List of clocks used in always blocks
        self.has_clock_crossing = False
        self.source_clock = None
        self.destination_clock = None


def extract_clocks_from_statements(statement, input_ports):
    """
    Recursively extracts clock signals used in statements inside always blocks.

    :param statement: The AST node representing the statement to process.
    :param input_ports: Set of input port names for the module.
    :return: Set of clock signal names found in the statements.
    """
    clocks = set()

    if isinstance(statement, IfStatement):
        # Process the condition of the IfStatement
        condition_signals = extract_signals_from_expression(statement.cond)
        # Identify clock signals used in the condition
        for sig in condition_signals:
            if sig in input_ports and is_clock_signal(sig):
                clocks.add(sig)
        # Recursively process the 'true_statement' and 'false_statement'
        clocks.update(extract_clocks_from_statements(statement.true_statement, input_ports))
        if statement.false_statement:
            clocks.update(extract_clocks_from_statements(statement.false_statement, input_ports))
    elif isinstance(statement, CaseStatement):
        # Process the expression used in the case statement
        case_expr_signals = extract_signals_from_expression(statement.comp)
        for sig in case_expr_signals:
            if sig in input_ports and is_clock_signal(sig):
                clocks.add(sig)
        # Recursively process each case item
        for case_item in statement.caselist:
            clocks.update(extract_clocks_from_statements(case_item.statement, input_ports))
    elif isinstance(statement, ForStatement):
        # Process the initialization, condition, and increment expressions
        if statement.init:
            clocks.update(extract_clocks_from_statements(statement.init, input_ports))
        if statement.cond:
            clocks.update(extract_clocks_from_statements(statement.cond, input_ports))
        if statement.next:
            clocks.update(extract_clocks_from_statements(statement.next, input_ports))
        # Process the body of the for loop
        clocks.update(extract_clocks_from_statements(statement.statement, input_ports))
    elif isinstance(statement, Block):
        # Process each statement in the block
        for stmt in statement.statements:
            clocks.update(extract_clocks_from_statements(stmt, input_ports))
    elif isinstance(statement, (Assign, NonblockingSubstitution)):
        # Process the right-hand side expression of the assignment
        rhs_signals = extract_signals_from_expression(statement.right)
        for sig in rhs_signals:
            if sig in input_ports and is_clock_signal(sig):
                clocks.add(sig)
    elif hasattr(statement, 'statement'):
        # Process nested statements in loops or other constructs
        clocks.update(extract_clocks_from_statements(statement.statement, input_ports))
    elif isinstance(statement, list):
        # Process a list of statements
        for stmt in statement:
            clocks.update(extract_clocks_from_statements(stmt, input_ports))
    else:
        # Handle other types of statements if needed
        pass

    return clocks


def extract_clock_domains_from_asts(asts):
    """
    Extracts clock domain information from the ASTs based on clocks used in always blocks.
    Returns a dictionary with module names as keys and their clock domain information as values.
    """
    # Build a mapping from module names to their AST definitions
    module_definitions = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                module_definitions[module_name] = definition

    # Extract input ports for each module
    input_ports_dict = extract_input_ports(module_definitions)

    clock_domains_info_dict = {}
    for module_name, definition in module_definitions.items():
        module_info = ModuleClockDomainInfo(module_name)
        # Process the module to extract clock domain information
        clocks_in_module = []
        clock_set = set()
        input_ports = input_ports_dict.get(module_name, set())
        for item in definition.items:
            if isinstance(item, Always):
                sens_list = item.sens_list
                # Extract the sensitivity list
                if sens_list is not None:
                    for sens in sens_list.list:
                        if hasattr(sens, 'sig'):
                            sig_name = get_rvalue(sens.sig)
                            # Include only clock signals, exclude reset signals
                            if is_clock_signal(sig_name) and sig_name not in clock_set:
                                clocks_in_module.append(sig_name)
                                clock_set.add(sig_name)
                # Extract clocks from statements inside the always block
                if item.statement is not None:
                    clocks_in_statements = extract_clocks_from_statements(item.statement, input_ports)
                    for clk in clocks_in_statements:
                        if clk not in clock_set:
                            clocks_in_module.append(clk)
                            clock_set.add(clk)
        # Store the clocks found in always blocks
        module_info.clocks_in_module = clocks_in_module
        # Check for clock crossings
        if len(clocks_in_module) > 1:
            module_info.has_clock_crossing = True
            # Assign source and destination clocks based on the order they appear
            module_info.source_clock = clocks_in_module[0]
            module_info.destination_clock = clocks_in_module[1]
        else:
            module_info.has_clock_crossing = False
            if module_info.clocks_in_module:
                module_info.source_clock = module_info.clocks_in_module[0]
            else:
                module_info.source_clock = None
        # Store the module info
        clock_domains_info_dict[module_name] = module_info
    return clock_domains_info_dict


def get_port_name_by_position(module_definitions, module_name, position):
    """
    Retrieves the port name from the module definition based on its position for positional connections.

    :param module_definitions: Dictionary mapping module names to their AST definitions.
    :param module_name: The name of the module whose port name is to be retrieved.
    :param position: The position index of the port in the module's port list.
    :return: The port name if found, otherwise a placeholder name.
    """
    module_def = module_definitions.get(module_name)
    if module_def and module_def.portlist and module_def.portlist.ports:
        if position < len(module_def.portlist.ports):
            port = module_def.portlist.ports[position]
            return port.name
    return f"port_{position}"


def is_reset_signal(signal_name):
    """
    Determines if a signal name corresponds to a reset signal based on common naming conventions.
    """
    reset_keywords = ['rst', 'reset', 'areset', 'sreset']
    signal_name_lower = signal_name.lower()
    return any(keyword in signal_name_lower for keyword in reset_keywords)


def print_clock_domains_info(clock_domains_info_dict):
    """
    Prints the clock domain information extracted from the RTL ASTs.
    """
    for module_name, module_info in clock_domains_info_dict.items():
        print(f"\nModule: {module_name}")
        if module_info.clocks_in_module:
            clocks = ', '.join(module_info.clocks_in_module)
            print(f"Clock: {clocks}")
        else:
            print("Clock: None")
        if module_info.has_clock_crossing:
            print(f"Clock Domains Crossing (CDC): source clock: {module_info.source_clock}, destination clock: {module_info.destination_clock}")


def print_clock_domains_check_info(clock_domains_info_dict, specs_clock_domains_text):
    """
    Prints the clock-domain information from both RTL and the specifications.

    :param clock_domains_info_dict: Dictionary from extract_clock_domains_from_asts(), keyed by module name,
                                    containing clock-domain details (clocks_in_module, has_clock_crossing, etc.)
    :param specs_clock_domains_text: The raw text (or processed text) from the specification's 'Clock Domains' chapter.
    """

    print("")
    print("RTL Clock Domains")
    print("=================")
    print_clock_domains_info(clock_domains_info_dict)

    print("")
    print("Specs Clock Domains")
    print("===================")
    print(textwrap.fill(specs_clock_domains_text, width=80))
    print("")


def compare_clock_domains(clock_domains_info_dict, specs_clock_domains_text):
    """
    Compares the clock domain information extracted from RTL with the specifications.
    Uses GPT to perform the comparison and outputs the result.
    """
    # Prepare the clock domain info text from RTL
    rtl_clock_domains_text = ''
    for module_name, module_info in clock_domains_info_dict.items():
        rtl_clock_domains_text += f"Module: {module_name}\n"
        clocks = ', '.join(module_info.clocks_in_module) if module_info.clocks_in_module else 'None'
        rtl_clock_domains_text += f"Clock: {clocks}\n\n"

    if DEBUG:
        print("RTL Clock Domains Text:\n", rtl_clock_domains_text)

    # Prepare the prompt for GPT
    prompt = (
        "You are an assistant that compares clock domain information from RTL code and specifications. "
        "You will be provided with the clock domain information extracted from RTL code and from specifications. "
        "Compare them based on module name, clock inputs, clock ports' name, clock domains, clock crossings, and any discrepancies in clock domain implementations. "
        "Pay close attention to numerical details and report any discrepancies. "
        "Do not consider reset signals as part of the clock domain information. "
        "If they match in all aspects, output:\n\n"
        "[INFO ] MATCHED\n\n"
        "If they do not match, output:\n\n"
        "[ERROR] MISMATCHED. <Provide the precise reason for the mismatch.>\n\n"
        "Only output the specified message, without any additional text.\n\n"
        "Here are the clock domain information:\n\n"
        "RTL Clock Domains Information:\n"
        "{}\n"
        "Specifications Clock Domains Information:\n"
        "{}"
    ).format(rtl_clock_domains_text.strip(), specs_clock_domains_text.strip())

    if DEBUG:
        print("GPT Prompt:\n", prompt)

    # Call GPT to perform the comparison
    try:
        completion = client.chat.completions.create(
            model='gpt-4o',
            messages=[
                {"role": "system",
                 "content": (
                     "You are an assistant that compares clock domain information from RTL code and specifications. "
                     "You only output the comparison result precisely as specified."
                 )},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=500,
            n=1,
            stop=None,
        )

        # Get the assistant's reply
        assistant_reply = completion.choices[0].message.content.strip()

        # Print the assistant's reply
        print(assistant_reply)

        # Return the result for final summary
        if "[INFO ]" in assistant_reply:
            result = assistant_reply.replace("[INFO ]", "").strip()
            return ("Clock Domains Check", result)
        elif "[ERROR]" in assistant_reply:
            result = assistant_reply.replace("[ERROR]", "").strip()
            return ("Clock Domains Check", result)
        else:
            return ("Clock Domains Check", assistant_reply)

    except Exception as e:
        print(f"[ERROR] Error comparing Clock Domains from specifications and RTL: {e}")
        return ("Clock Domains Check", "Error during comparison.")


class ModuleResetDomainInfo:
    def __init__(self, module_name):
        self.module_name = module_name
        self.resets_in_module = []    # List of resets used in always blocks
        self.has_reset_crossing = False
        self.source_reset = None
        self.destination_reset = None


def extract_input_ports(module_definitions):
    input_ports_dict = {}
    for module_name, definition in module_definitions.items():
        input_ports = set()

        # Process ports from definition.portlist
        if definition.portlist:
            for port in definition.portlist.ports:
                if isinstance(port, Ioport):
                    decl = port.first
                    if isinstance(decl, Input):
                        port_name = decl.name
                        input_ports.add(port_name)
                elif isinstance(port, Port):
                    port_name = port.name
                    direction = find_port_direction(definition.items, port_name)
                    if direction == 'input':
                        input_ports.add(port_name)

        # Process items in definition.items
        for item in definition.items:
            if isinstance(item, Decl):
                for decl in item.list:
                    if isinstance(decl, Input):
                        port_name = decl.name
                        input_ports.add(port_name)

        if DEBUG:
            print(f"Module: {module_name}, Input Ports: {input_ports}")

        input_ports_dict[module_name] = input_ports
    return input_ports_dict


def find_port_direction(items, port_name):
    """
    Finds the direction of a port by searching through the declarations in the module items.

    :param items: List of module items.
    :param port_name: The name of the port.
    :return: The direction of the port ('input', 'output', 'inout') or None if not found.
    """
    for item in items:
        if isinstance(item, Decl):
            for decl in item.list:
                if decl.name == port_name:
                    if isinstance(decl, Input):
                        return 'input'
                    elif isinstance(decl, Output):
                        return 'output'
                    elif isinstance(decl, Inout):
                        return 'inout'
    return None


def map_resets_to_clocks(definition, resets_in_module, input_ports):
    """
    Maps resets to their associated clocks by analyzing always blocks.

    :param definition: The module definition from the AST.
    :param resets_in_module: A set of reset signals found in the module.
    :param input_ports: A set of input port names for the module.
    :return: A dictionary mapping clock signals to sets of reset signals.
    """
    reset_clock_map = {}
    for item in definition.items:
        if isinstance(item, Always):
            sens_list = item.sens_list
            if sens_list:
                clocks = []
                resets = []
                for sens in sens_list.list:
                    if hasattr(sens, 'sig'):
                        sig_name = get_rvalue(sens.sig)
                        if sig_name in input_ports:
                            if is_clock_signal(sig_name):
                                clocks.append(sig_name)
                            elif is_reset_signal(sig_name):
                                resets.append(sig_name)
                # Map resets to clocks
                for clk in clocks:
                    for rst in resets:
                        if rst in resets_in_module:
                            if clk not in reset_clock_map:
                                reset_clock_map[clk] = set()
                            reset_clock_map[clk].add(rst)
    return reset_clock_map


def extract_reset_domains_from_asts(asts, clock_domains_info_dict):
    """
    Extracts reset domain information from the ASTs based on resets used in always blocks.
    Returns a dictionary with module names as keys and their reset domain information as values.
    """
    # Build a mapping from module names to their AST definitions
    module_definitions = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                module_definitions[module_name] = definition

    # Extract input ports for each module
    input_ports_dict = extract_input_ports(module_definitions)

    reset_domain_info_dict = {}
    for module_name, definition in module_definitions.items():
        module_info = ModuleResetDomainInfo(module_name)
        resets_in_module = set()
        input_ports = input_ports_dict.get(module_name, set())
        clock_info = clock_domains_info_dict.get(module_name, None)

        # Process the module to extract reset domain information
        for item in definition.items:
            if isinstance(item, Always):
                sens_list = item.sens_list
                # Extract resets from sensitivity list (asynchronous resets)
                if sens_list is not None:
                    for sens in sens_list.list:
                        if hasattr(sens, 'sig'):
                            sig_name = get_rvalue(sens.sig)
                            if sig_name in input_ports and is_reset_signal(sig_name):
                                resets_in_module.add(sig_name)
                # Process the statements inside the always block to find synchronous resets
                resets_in_module.update(extract_resets_from_statements(item.statement, input_ports))

        # Map resets to their associated clocks
        reset_clock_map = map_resets_to_clocks(definition, resets_in_module, input_ports)

        module_info.resets_in_module = list(resets_in_module)

        # Determine source and destination resets based on clock domains
        if len(reset_clock_map) > 1 and clock_info and clock_info.has_clock_crossing:
            module_info.has_reset_crossing = True
            source_clock = clock_info.source_clock
            destination_clock = clock_info.destination_clock
            source_resets = reset_clock_map.get(source_clock, set())
            destination_resets = reset_clock_map.get(destination_clock, set())
            module_info.source_reset = next(iter(source_resets), None)
            module_info.destination_reset = next(iter(destination_resets), None)
        else:
            module_info.has_reset_crossing = False
            if module_info.resets_in_module:
                module_info.source_reset = next(iter(module_info.resets_in_module), None)
            else:
                module_info.source_reset = None

        reset_domain_info_dict[module_name] = module_info

    return reset_domain_info_dict


def extract_resets_from_sens_list(sens_list, input_ports):
    """
    Extracts reset signals from the sensitivity list of an always block.

    :param sens_list: Sensitivity list of the always block.
    :param input_ports: Set of input port names for the module.
    :return: Set of reset signal names.
    """
    resets = set()
    clock_signals = set()
    for sens in sens_list.list:
        if hasattr(sens, 'sig'):
            sig_name = get_rvalue(sens.sig)
            if sig_name in input_ports:
                if is_clock_signal(sig_name):
                    clock_signals.add(sig_name)
                elif is_reset_signal(sig_name):
                    resets.add(sig_name)
    # An asynchronous reset is typically in the sensitivity list along with a clock signal
    if resets:
        return resets
    else:
        return set()


def extract_resets_from_statements(statement, input_ports):
    """
    Recursively extracts reset signals used in conditional statements inside always blocks.

    :param statement: The AST node representing the statement to process.
    :param input_ports: Set of input port names for the module.
    :return: Set of reset signal names found in the statements.
    """
    resets = set()

    if isinstance(statement, IfStatement):
        # Process the condition of the IfStatement
        condition_signals = extract_signals_from_expression(statement.cond)
        # Identify reset signals used in the condition
        for sig in condition_signals:
            if sig in input_ports and is_reset_signal(sig):
                resets.add(sig)
                if DEBUG:
                    print(f"Found reset signal '{sig}' in IfStatement condition.")
        # Recursively process the 'true_statement' and 'false_statement'
        resets.update(extract_resets_from_statements(statement.true_statement, input_ports))
        if statement.false_statement:
            resets.update(extract_resets_from_statements(statement.false_statement, input_ports))
    elif isinstance(statement, CaseStatement):
        # Process the expression used in the case statement
        case_expr_signals = extract_signals_from_expression(statement.comp)
        for sig in case_expr_signals:
            if sig in input_ports and is_reset_signal(sig):
                resets.add(sig)
                if DEBUG:
                    print(f"Found reset signal '{sig}' in CaseStatement expression.")
        # Recursively process each case item
        for case_item in statement.caselist:
            resets.update(extract_resets_from_statements(case_item.statement, input_ports))
    elif isinstance(statement, ForStatement):
        # Process the initialization, condition, and increment expressions
        if statement.init:
            resets.update(extract_resets_from_statements(statement.init, input_ports))
        if statement.cond:
            resets.update(extract_resets_from_statements(statement.cond, input_ports))
        if statement.next:
            resets.update(extract_resets_from_statements(statement.next, input_ports))
        # Process the body of the for loop
        resets.update(extract_resets_from_statements(statement.statement, input_ports))
    elif isinstance(statement, Block):
        # Process each statement in the block
        for stmt in statement.statements:
            resets.update(extract_resets_from_statements(stmt, input_ports))
    elif isinstance(statement, (Assign, NonblockingSubstitution)):
        # Process the right-hand side expression of the assignment
        rhs_signals = extract_signals_from_expression(statement.right)
        for sig in rhs_signals:
            if sig in input_ports and is_reset_signal(sig):
                resets.add(sig)
                if DEBUG:
                    print(f"Found reset signal '{sig}' in assignment.")
    elif hasattr(statement, 'statement'):
        # Process nested statements in loops or other constructs
        resets.update(extract_resets_from_statements(statement.statement, input_ports))
    elif isinstance(statement, list):
        # Process a list of statements
        for stmt in statement:
            resets.update(extract_resets_from_statements(stmt, input_ports))
    else:
        # Handle other types of statements if needed
        pass

    return resets


def extract_signals_from_expression(expr):
    """
    Recursively extracts signal names from an expression.
    """
    signals = set()
    if isinstance(expr, Identifier):
        signals.add(expr.name)
    elif hasattr(expr, 'children'):
        for child in expr.children():
            signals.update(extract_signals_from_expression(child))
    return signals


def print_reset_domain_info(reset_domain_info_dict):
    """
    Prints the reset domain information extracted from the RTL ASTs.
    """
    for module_name, module_info in reset_domain_info_dict.items():
        print(f"\nModule: {module_name}")
        if module_info.resets_in_module:
            resets = ', '.join(module_info.resets_in_module)
            print(f"Reset: {resets}")
        else:
            print("Reset: None")
        if module_info.has_reset_crossing:
            print(f"Reset Domain Crossing (RDC): source reset: {module_info.source_reset}, destination reset: {module_info.destination_reset}")


def print_reset_domains_check_info(reset_domain_info_dict, specs_reset_domains_text):
    """
    Prints the reset-domain information from both RTL and the specifications.

    :param reset_domain_info_dict: Dictionary from extract_reset_domains_from_asts(), keyed by module name,
                                   containing reset-domain details (resets_in_module, has_reset_crossing, etc.)
    :param specs_reset_domains_text: The raw text (or processed text) from the specification's 'Reset Domains' chapter.
    """

    print("")
    print("RTL Reset Domains")
    print("=================")
    print_reset_domain_info(reset_domain_info_dict)

    print("")
    print("Specs Reset Domains")
    print("===================")
    # Optionally wrap the spec text to 80 columns, or any column width
    print(textwrap.fill(specs_reset_domains_text, width=80))
    print("")


def compare_reset_domains(reset_domain_info_dict, specs_reset_domain_text):
    """
    Compares the reset domain information extracted from RTL with the specifications.
    Uses GPT to perform the comparison and outputs the result.
    """
    # Prepare the reset domain info text from RTL
    rtl_reset_domain_text = ''
    for module_name, module_info in reset_domain_info_dict.items():
        rtl_reset_domain_text += f"Module: {module_name}\n"
        resets = ', '.join(module_info.resets_in_module) if module_info.resets_in_module else 'None'
        rtl_reset_domain_text += f"Reset: {resets}\n\n"

    if DEBUG:
        print("RTL Reset Domains Text:\n", rtl_reset_domain_text)

    # Prepare the prompt for GPT
    prompt = (
        "You are an assistant that compares reset domain information from RTL code and specifications. "
        "You will be provided with the reset domains information extracted from RTL code and from specifications. "
        "Compare them based on module name, reset inputs, reset ports' name, reset domains, reset crossings, and any discrepancies in reset domain implementations. "
        "Pay close attention to details and report any discrepancies. "        
        "If they match in all aspects, output:\n\n"
        "[INFO ] MATCHED\n\n"
        "If they do not match, output:\n\n"
        "[ERROR] MISMATCHED. <Provide the precise reason for the mismatch.>\n\n"
        "Only output the specified message, without any additional text.\n\n"
        "Here are the reset domains information:\n\n"
        "RTL Reset Domains Information:\n"
        "{}\n"
        "Specifications Reset Domains Information:\n"
        "{}"
    ).format(rtl_reset_domain_text.strip(), specs_reset_domain_text.strip())

    if DEBUG:
        print("GPT Prompt:\n", prompt)

    # Call GPT to perform the comparison
    try:
        completion = client.chat.completions.create(
            model='gpt-4o',
            messages=[
                {"role": "system",
                 "content": (
                     "You are an assistant that compares reset domains information from RTL code and specifications. "
                     "You only output the comparison result precisely as specified."
                 )},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=500,
            n=1,
            stop=None,
        )

        # Get the assistant's reply
        assistant_reply = completion.choices[0].message.content.strip()

        # Print the assistant's reply
        print(assistant_reply)

        # Return the result for final summary
        if "[INFO ]" in assistant_reply:
            result = assistant_reply.replace("[INFO ]", "").strip()
            return ("Reset Domains Check", result)
        elif "[ERROR]" in assistant_reply:
            result = assistant_reply.replace("[ERROR]", "").strip()
            return ("Reset Domains Check", result)
        else:
            return ("Reset Domains Check", assistant_reply)

    except Exception as e:
        print(f"[ERROR] Error comparing Reset Domains from specifications and RTL: {e}")
        return ("Reset Domains Check", "Error during comparison.")


class StateMachine:
    def __init__(self, module_name):
        self.module_name = module_name
        self.state_var = None
        self.next_state_var = None
        self.states = {}  # Mapping from state values to state names
        self.transitions = []  # List of transitions, each as (current_state, condition, next_state)


def extract_state_machines_from_asts(asts):
    state_machines = {}
    for ast in asts:
        for description in ast.description.definitions:
            if hasattr(description, 'name'):
                module_name = description.name
                state_machine = StateMachine(module_name)
                module_definitions = description.items
                # Extract state variable names and state parameters
                extract_state_variables_and_parameters(module_definitions, state_machine)
                # Find state update and state transition always blocks
                find_state_always_blocks(module_definitions, state_machine)
                if state_machine.state_var and state_machine.transitions:
                    state_machines[module_name] = state_machine
    return state_machines

def extract_state_variables_and_parameters(module_items, state_machine):
    for item in module_items:
        if isinstance(item, Decl):
            for decl in item.list:
                if isinstance(decl, Reg):
                    if 'state' in decl.name:
                        if not state_machine.state_var:
                            state_machine.state_var = decl.name
                            print(f"Found state variable: {state_machine.state_var}")
                        else:
                            state_machine.next_state_var = decl.name
                            print(f"Found next state variable: {state_machine.next_state_var}")
                elif isinstance(decl, (Localparam, Parameter)):
                    # Extract state values
                    state_name = decl.name
                    state_value = get_rvalue(decl.value)
                    state_machine.states[state_value] = state_name
                    print(f"Found state constant: {state_name} = {state_value}")


def find_state_always_blocks(module_items, state_machine):
    for item in module_items:
        if isinstance(item, Always):
            print(f"Always block statement type: {type(item.statement)}")
            print_ast_node(item.statement, "")
            assigned_vars = get_assigned_vars(item.statement)
            print(f"Assigned vars in always block: {assigned_vars}")
            if state_machine.state_var in assigned_vars:
                # This is the state register update always block
                continue
            elif state_machine.next_state_var in assigned_vars:
                # This is the state transition always block
                print("Found state transition always block")
                extract_state_transitions(item.statement, state_machine)


def get_assigned_vars(statement):
    assigned_vars = set()
    if isinstance(statement, Block):
        for stmt in statement.statements:
            assigned_vars.update(get_assigned_vars(stmt))
    elif isinstance(statement, IfStatement):
        assigned_vars.update(get_assigned_vars(statement.true_statement))
        if statement.false_statement:
            assigned_vars.update(get_assigned_vars(statement.false_statement))
    elif isinstance(statement, CaseStatement):
        for case_item in statement.caselist:
            case_stmt = case_item.statement
            if isinstance(case_stmt, (list, tuple)):
                for stmt in case_stmt:
                    assigned_vars.update(get_assigned_vars(stmt))
            else:
                assigned_vars.update(get_assigned_vars(case_stmt))
    elif isinstance(statement, (NonblockingSubstitution, BlockingSubstitution, Assign)):
        # Extract identifiers from the left-hand side recursively
        left_vars = extract_identifiers(statement.left)
        assigned_vars.update(left_vars)
    elif hasattr(statement, 'statement'):
        assigned_vars.update(get_assigned_vars(statement.statement))
    elif isinstance(statement, (list, tuple)):
        for stmt in statement:
            assigned_vars.update(get_assigned_vars(stmt))
    else:
        # Handle other types if necessary
        pass
    return assigned_vars

def extract_identifiers(node):
    identifiers = set()
    if isinstance(node, Identifier):
        identifiers.add(node.name)
    elif hasattr(node, 'children'):
        for child in node.children():
            identifiers.update(extract_identifiers(child))
    return identifiers


def extract_state_transitions(statement, state_machine, current_state=None):
    condition = 'default'  # Initialize condition
    if isinstance(statement, CaseStatement):
        # Extract the expression being switched on
        if isinstance(statement.comp, Identifier):
            if statement.comp.name == state_machine.state_var:
                # Top-level case statement on the state variable
                for case_item in statement.caselist:
                    # Get the case condition(s)
                    case_conditions = case_item.cond
                    if not isinstance(case_conditions, list):
                        case_conditions = [case_conditions]
                    for case_cond in case_conditions:
                        state_value = get_rvalue(case_cond)
                        state_name = state_machine.states.get(state_value, state_value)
                        # Recursively extract transitions for this state
                        extract_state_transitions(case_item.statement, state_machine, current_state=state_name)
    elif isinstance(statement, IfStatement):
        condition = get_condition_str(statement.cond)
        # Extract transitions from true_statement
        extract_state_transitions(statement.true_statement, state_machine, current_state)
        # Extract transitions from false_statement
        if statement.false_statement:
            extract_state_transitions(statement.false_statement, state_machine, current_state)
    elif isinstance(statement, (NonblockingSubstitution, BlockingSubstitution)):
        if isinstance(statement.left, Identifier) and statement.left.name == state_machine.next_state_var:
            next_state_value = get_rvalue(statement.right)
            next_state_name = state_machine.states.get(next_state_value, next_state_value)
            state_machine.transitions.append((current_state, condition, next_state_name))
            print(f"Added transition: {current_state} --({condition})--> {next_state_name}")
    elif isinstance(statement, Block):
        for stmt in statement.statements:
            extract_state_transitions(stmt, state_machine, current_state)
    elif isinstance(statement, list):
        for stmt in statement:
            extract_state_transitions(stmt, state_machine, current_state)
    elif hasattr(statement, 'statement'):
        extract_state_transitions(statement.statement, state_machine, current_state)
    else:
        # Handle other statement types if necessary
        pass


def get_condition_str(condition):
    """
    Converts a condition expression into a string representation.

    :param condition: The AST node representing the condition.
    :return: A string representation of the condition.
    """
    return get_rvalue(condition)


def get_assigned_states(statement, var_name):
    next_states = []
    if isinstance(statement, (NonblockingSubstitution, BlockingSubstitution)):
        if isinstance(statement.left, Identifier) and statement.left.name == var_name:
            next_state_value = get_rvalue(statement.right)
            next_states.append(next_state_value)
    elif isinstance(statement, Block):
        for stmt in statement.statements:
            next_states.extend(get_assigned_states(stmt, var_name))
    elif isinstance(statement, IfStatement):
        next_states.extend(get_assigned_states(statement.true_statement, var_name))
        if statement.false_statement:
            next_states.extend(get_assigned_states(statement.false_statement, var_name))
    elif isinstance(statement, CaseStatement):
        for case_item in statement.caselist:
            next_states.extend(get_assigned_states(case_item.statement, var_name))
    return next_states


def print_state_machine_info(state_machines):
    for module_name, state_machine in state_machines.items():
        print(f"\nState Machine in Module: {module_name}")
        print(f"State Variable: {state_machine.state_var}")
        print(f"Next State Variable: {state_machine.next_state_var}")
        print("\nStates:")
        for state_value, state_name in state_machine.states.items():
            print(f"  {state_name} ({state_value})")
        print("\nTransitions:")
        for current_state, condition, next_state in state_machine.transitions:
            current_state_name = state_machine.states.get(current_state, current_state)
            next_state_name = state_machine.states.get(next_state, next_state)
            print(f"  From {current_state_name} -> {next_state_name} on {condition}")
        print("\n")


def compare_state_machines(rtl_state_machine, specs_chapter_content, module_name):
    """
    Compares the state machine extracted from RTL with the specifications.
    Uses GPT to perform the comparison and outputs the result.
    """
    # Prepare the RTL state machine description
    rtl_state_machine_description = generate_rtl_state_machine_description(rtl_state_machine)

    # Prepare the prompt
    prompt = (
        "You are an assistant that compares state machines from RTL code and specifications. "
        "You will be provided with a state machine extracted from RTL code and a state machine description from specifications. "
        "Compare them based on state names, transitions, and conditions. "
        "Pay close attention to numerical details and report any discrepancies. "
        "If they match in all aspects, output:\n\n"
        "[INFO ] MATCHED\n\n"
        "If they do not match, output:\n\n"
        "[ERROR] MISMATCHED. <Provide the precise reason for the mismatch, including specific differences in state names, transitions, or conditions.>\n\n"
        "Only output the specified message, without any additional text.\n\n"
        "Here are the state machines for module {}:\n\n"
        "RTL State Machine:\n"
        "{}\n"
        "Specifications State Machine:\n"
        "{}"
    ).format(module_name, rtl_state_machine_description.strip(), specs_chapter_content.strip())

    if DEBUG:
        print("GPT Prompt:\n", prompt)

    # Call GPT to perform the comparison
    try:
        completion = client.chat.completions.create(
            model='gpt-4o',
            messages=[
                {"role": "system",
                 "content": (
                     "You are an assistant that compares state machines from RTL code and specifications. "
                     "You only output the comparison result precisely as specified."
                 )},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=500,
            n=1,
            stop=None,
        )

        # Get the assistant's reply
        assistant_reply = completion.choices[0].message.content.strip()

        # Print the assistant's reply
        print(assistant_reply)

        # Return the result for final summary
        if "[INFO ]" in assistant_reply:
            result = assistant_reply.replace("[INFO ]", "").strip()
            return ("State Machine Check", result)
        elif "[ERROR]" in assistant_reply:
            result = assistant_reply.replace("[ERROR]", "").strip()
            return ("State Machine Check", result)
        else:
            return ("State Machine Check", assistant_reply)

    except Exception as e:
        print(f"[ERROR] Error comparing state machines from specifications and RTL: {e}")
        return ("State Machine Check", "Error during comparison.")


def generate_rtl_state_machine_description(state_machine):
    """
    Generates a textual description of the RTL state machine for inclusion in the GPT prompt.
    """
    description = f"State Variable: {state_machine.state_var}\n"
    description += f"Next State Variable: {state_machine.next_state_var}\n\n"
    description += "States:\n"
    for state_value, state_name in state_machine.states.items():
        description += f"  {state_name} ({state_value})\n"
    description += "\nTransitions:\n"
    for current_state, condition, next_state in state_machine.transitions:
        current_state_name = state_machine.states.get(current_state, current_state)
        next_state_name = state_machine.states.get(next_state, next_state)
        description += f"  From {current_state_name} -> {next_state_name} on {condition}\n"
    return description


def print_single_state_machine_info(state_machine):
    if state_machine is None:
        print("No state machine information available.")
        return

    print(f"\nState Machine in Module: {state_machine.module_name}")
    print(f"State Variable: {state_machine.state_var}")
    print(f"Next State Variable: {state_machine.next_state_var}")
    print("\nStates:")
    for state_value, state_name in state_machine.states.items():
        print(f"  {state_name} ({state_value})")
    print("\nTransitions:")
    for current_state, condition, next_state in state_machine.transitions:
        current_state_name = state_machine.states.get(current_state, current_state)
        next_state_name = state_machine.states.get(next_state, next_state)
        print(f"  From {current_state_name} -> {next_state_name} on {condition}")
    print("\n")

def main():
    run_rtl_parameters_check = True
    run_io_ports_check = True
    run_module_hierarchy_check = True
    run_clock_domains_check = True
    run_reset_domains_check = True

    # Path to your specifications Word document
    spec_docx_path = 'specs/specs.docx'

    # # List of modules' name
    module_names = ['ufifo', 'rxuart', 'txuart', 'wbuart']

    # Get the list of Verilog files
    file_list = get_file_list("rtl")

    # Generate ASTs for the files
    asts = generate_asts(file_list)

    # Prints the AST contents for each module.
    if DEBUG:
        print_ast_contents(asts)

    # Extract chapters from the Word document
    chapters = extract_text_by_chapters(spec_docx_path)
    final_results = []  # Store the results as tuples (check_name, result_content)

    # Find the chapter with title 'RTL Parameters'
    parameters_chapter = next((chapter for chapter in chapters if chapter['title'] == 'RTL Parameters'), None)

    if run_rtl_parameters_check:
        if parameters_chapter:
            # Add header for RTL Default Parameters check
            print("\n===== RTL Default Parameters Check =====\n")

            # Extract parameters from specifications for each module
            specs_params_dict = {module_name: extract_parameters_from_specs(parameters_chapter, module_name) for module_name in module_names}

            # Extract parameters from each AST
            rtl_params_dict = extract_parameters_from_asts(asts)

            # Print the extracted parameters from RTL
            if DEBUG:
                print_parameters(rtl_params_dict)

            # Compare parameters
            final_results.append(compare_parameters(rtl_params_dict, specs_params_dict))
        else:
            print("[ERROR] Chapter 'RTL Parameters' not found in the specifications.")

    # Find the chapter with title 'IO Ports'
    ioports_chapter = next((chapter for chapter in chapters if chapter['title'] == 'IO Ports'), None)

    if run_io_ports_check:
        if ioports_chapter:
            # Add header for IO Ports check
            print("\n===== IO Ports Check =====\n")

            # Extract IO ports' information from specifications for each module
            spec_ioports_dict = {module_name: extract_ioports_from_specs(ioports_chapter, module_name) for module_name in module_names}

            # Extract IO ports' information from each AST
            rtl_ioports_dict = extract_ioports_from_asts(asts)

            # Print the extracted IO ports' information from RTL
            if DEBUG:
                print_ioports(rtl_ioports_dict)

            # Compare IO ports' information
            final_results.append(compare_ioports(rtl_ioports_dict, spec_ioports_dict))
        else:
            print("[ERROR] Chapter 'IO Ports' not found in the specifications.")

    # Find the chapter with title 'Module Hierarchy'
    module_hierarchy_chapter = next((chapter for chapter in chapters if chapter['title'] == 'Module Hierarchy'), None)

    if run_module_hierarchy_check:
        if module_hierarchy_chapter:
            # Add header for Module Hierarchy check
            print("\n===== Module Hierarchy Check =====\n")

            specs_module_hierarchy_text = module_hierarchy_chapter['content']

            # Extract module hierarchy
            hierarchy_trees = extract_module_hierarchy(asts)

            # Print the module hierarchy
            if DEBUG:
                print_module_hierarchy(hierarchy_trees)

            # Compare Module Hierarchy
            final_results.append(compare_module_hierarchy(hierarchy_trees, specs_module_hierarchy_text))
        else:
            print("[ERROR] Chapter 'Module Hierarchy' not found in the specifications.")

    # Find the chapter with title 'Clock Domains' in the specifications
    clock_domains_chapter = next((chapter for chapter in chapters if chapter['title'] == 'Clock Domains'), None)

    if run_clock_domains_check:
        if clock_domains_chapter:
            # Add header for Clock Domains Check
            print("\n===== Clock Domains Check =====\n")

            # Extract clock domains information from ASTs
            clock_domains_info_dict = extract_clock_domains_from_asts(asts)

            specs_clock_domains_text = clock_domains_chapter['content']

            # Print the extracted clock domains information
            print_clock_domains_check_info(clock_domains_info_dict, specs_clock_domains_text)

            # Compare Clock Domains
            final_results.append(compare_clock_domains(clock_domains_info_dict, specs_clock_domains_text))
        else:
            print("[ERROR] Chapter 'Clock Domains' not found in the specifications.")

    # Find the chapter with title 'Reset Domains' in the specifications
    reset_domains_chapter = next((chapter for chapter in chapters if chapter['title'] == 'Reset Domains'), None)

    if run_reset_domains_check:
        if reset_domains_chapter:
            # Add header for Reset Domains Check
            print("\n===== Reset Domains Check =====\n")

            # Extract clock domains information from ASTs
            clock_domains_info_dict = extract_clock_domains_from_asts(asts)

            # Extract reset domains information from ASTs, passing clock domain info
            reset_domains_info_dict = extract_reset_domains_from_asts(asts, clock_domains_info_dict)

            specs_reset_domains_text = reset_domains_chapter['content']

            # Print the extracted reset domains information
            print_reset_domains_check_info(reset_domains_info_dict, specs_reset_domains_text)

            # Compare Reset Domains
            final_results.append(compare_reset_domains(reset_domains_info_dict, specs_reset_domains_text))
        else:
            print("[ERROR] Chapter 'Reset Domains' not found in the specifications.")

    # Calculate the maximum length of check names for alignment
    if final_results:
        max_check_name_length = max(len(check[0]) for check in final_results)

        # Print all results with aligned check names
        print("\n===== Final Results =====\n")
        for check_name, result_content in final_results:
            print(f"{check_name.ljust(max_check_name_length)} : {result_content}")
        print("\n=========================")

if __name__ == "__main__":
    main()

