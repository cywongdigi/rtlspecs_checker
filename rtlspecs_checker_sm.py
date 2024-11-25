import os
import base64
import olefile
import io
import openai
import csv
import docx
import sympy
from pyverilog.vparser.parser import parse
from pyverilog.vparser.ast import (
    Node, Decl, Reg, Wire, Localparam, Parameter, IntConst, StringConst, Identifier,
    Partselect, Pointer, NonblockingSubstitution,
    BlockingSubstitution, IfStatement, CaseStatement, Block, Repeat,
    Concat, LConcat, CasezStatement, Case, Assign, Always,
    InstanceList, Instance, Ioport, Input, Output, Inout, Port,
    Rvalue, Lvalue, Cond, Width, # Add Rvalue, Lvalue, Cond
    # Operator classes
    Plus, Minus, Times, Divide, Mod, Power, Eq, Eql, NotEq, NotEql,
    GreaterThan, GreaterEq, LessThan, LessEq, And, Or, Xor, Xnor,
    Land, Lor, Sll, Srl, Sra, Uplus, Uminus, Ulnot, Unot,
    Uand, Unand, Uor, Unor, Uxor, Uxnor,
)
from docx import Document
from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image
from openai import OpenAI
import traceback  # For exception handling
from collections import defaultdict, deque
from sympy import symbols, S
from sympy.logic.boolalg import Boolean, BooleanFunction, And as SymAnd, Or as SymOr, Not as SymNot

# Load your OpenAI API key from the environment variable
openai.api_key = os.getenv('OPENAI_API_KEY')

# Set your OpenAI API key
client = openai

# Set to True to enable debug output
DEBUG = False


class ModuleNode:
    def __init__(self, module_name, instance_name=None):
        self.module_name = module_name
        self.instance_name = instance_name
        self.children = []  # List of ModuleNode instances

    def add_child(self, child_node):
        self.children.append(child_node)

    def print_hierarchy(self, level=0):
        indent = "    " * level
        if self.instance_name:
            print(f"{indent}- Instance: {self.instance_name}, Module: {self.module_name}")
        else:
            print(f"Top-Level Module: {self.module_name}")
        for child in self.children:
            child.print_hierarchy(level + 1)


def get_file_list(directory, extension=".v"):
    """
    Returns a list of Verilog files in the specified directory.

    :param directory: Directory to search for Verilog files.
    :param extension: File extension to filter by (default is ".v").
    :return: List of Verilog file paths.
    """
    return [os.path.join(directory, f) for f in os.listdir(directory) if f.endswith(extension)]


def generate_asts(file_list):
    """
    Parses Verilog files and returns their ASTs.

    :param file_list: List of Verilog file paths.
    :return: List of ASTs corresponding to the Verilog files.
    """
    asts = []
    for filename in file_list:
        try:
            ast, _ = parse([filename])
            asts.append(ast)
        except Exception as e:
            print(f"[ERROR] Error parsing {filename}: {e}")
    return asts


def print_ast_contents(asts):
    """
    Prints the AST contents for each module.

    :param asts: List of ASTs corresponding to the Verilog files.
    """
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                print(f"\nAST for module: {module_name}")
                print_ast_node(definition, "")


def print_ast_node(node, indent):
    if isinstance(node, Node):
        node_info = f"{indent}{node.__class__.__name__}"

        # Add more detailed information for specific types of nodes
        if isinstance(node, (Parameter, Localparam)):
            node_info += f" (name={node.name}, value={get_rvalue(node.value)})"
        elif isinstance(node, Reg):
            node_info += f" (name={node.name}, width={get_rvalue(node.width)})"
        elif isinstance(node, Assign):
            node_info += f" (LHS={get_lvalue(node.left)}, RHS={get_rvalue(node.right)})"
        elif isinstance(node, Always):
            sens_list = node.sens_list if node.sens_list else 'None'
            node_info += f" (sensitivity={sens_list})"
        elif isinstance(node, IfStatement):
            condition_str = get_rvalue(node.cond)
            node_info += f" (condition={condition_str})"
        elif isinstance(node, CaseStatement):
            comp_expr = get_rvalue(node.comp)
            node_info += f" (expression={comp_expr})"
        elif isinstance(node, (Plus, Minus, Times, Divide, Mod, Power,
                               Eq, Eql, NotEq, NotEql, GreaterThan, GreaterEq,
                               LessThan, LessEq, And, Or, Xor, Xnor,
                               Land, Lor, Sll, Srl, Sra, Uand, Unand, Uor, Unor, Uxor, Uxnor)):
            node_info += f" (operator={node.__class__.__name__})"
        else:
            if hasattr(node, 'name'):
                node_info += f" (name={node.name})"
            if hasattr(node, 'value'):
                node_info += f" (value={get_rvalue(node.value)})"

        print(node_info)

        # Recursively print each child node with increased indentation
        for child in node.children():
            print_ast_node(child, indent + "    ")
    else:
        # Print non-Node types directly (e.g., strings, integers)
        print(f"{indent}{node}")


def get_lvalue(node):
    if isinstance(node, Lvalue):
        return get_lvalue(node.var)  # Lvalue wraps another node in .var
    elif isinstance(node, Identifier):
        return node.name
    elif isinstance(node, Pointer):
        var = get_lvalue(node.var)
        ptr = get_rvalue(node.ptr)
        return f"{var}[{ptr}]"
    elif isinstance(node, Partselect):
        var = get_lvalue(node.var)
        msb = get_rvalue(node.msb)
        lsb = get_rvalue(node.lsb)
        return f"{var}[{msb}:{lsb}]"
    else:
        # Handle other types if necessary
        return str(node)


def get_rvalue(value):
    if isinstance(value, Rvalue):
        return get_rvalue(value.var)
    elif isinstance(value, IntConst):
        return value.value
    elif isinstance(value, StringConst):
        return value.value
    elif isinstance(value, Identifier):
        return value.name
    elif isinstance(value, (Uplus, Uminus, Ulnot, Unot, Uand, Unand, Uor, Unor, Uxor, Uxnor)):
        # Unary operators handling
        operator = get_operator_symbol(value.__class__.__name__)
        operand = get_rvalue(value.right if hasattr(value, 'right') else value.var)
        return f"({operator}{operand})"
    elif isinstance(value, (Plus, Minus, Times, Divide, Mod, Power,
                            Eq, Eql, NotEq, NotEql, GreaterThan, GreaterEq,
                            LessThan, LessEq, And, Or, Xor, Xnor,
                            Land, Lor, Sll, Srl, Sra)):
        # Binary operators handling
        operator = get_operator_symbol(value.__class__.__name__)
        left = get_rvalue(value.left)
        right = get_rvalue(value.right)
        return f"({left} {operator} {right})"
    elif isinstance(value, Partselect):
        var = get_rvalue(value.var)
        msb = get_rvalue(value.msb)
        lsb = get_rvalue(value.lsb)
        return f"{var}[{msb}:{lsb}]"
    elif isinstance(value, Pointer):
        var = get_rvalue(value.var)
        ptr = get_rvalue(value.ptr)
        return f"{var}[{ptr}]"
    elif isinstance(value, Concat):
        items = [get_rvalue(child) for child in value.children()]
        return f"{{{', '.join(items)}}}"
    elif isinstance(value, Repeat):
        times = get_rvalue(value.times)
        value_r = get_rvalue(value.value)
        return f"{{ {times}{{{value_r}}} }}"
    elif isinstance(value, Cond):
        cond = get_rvalue(value.cond)
        true_value = get_rvalue(value.true_value)
        false_value = get_rvalue(value.false_value)
        return f"({cond}) ? ({true_value}) : ({false_value})"
    elif isinstance(value, Width):
        msb = get_rvalue(value.msb)
        lsb = get_rvalue(value.lsb)
        return f"[{msb}:{lsb}]"
    elif isinstance(value, (list, tuple)):
        # Convert list or tuple to string
        return ', '.join(get_rvalue(v) for v in value)
    elif value is None:
        return ''
    else:
        return str(value)




def get_operator_symbol(op_class_name):
    operator_map = {
        'Plus': '+',
        'Minus': '-',
        'Times': '*',
        'Divide': '/',
        'Mod': '%',
        'Power': '**',
        'Eq': '==',
        'Eql': '===',
        'NotEq': '!=',
        'NotEql': '!==',
        'GreaterThan': '>',
        'GreaterEq': '>=',
        'LessThan': '<',
        'LessEq': '<=',
        'And': '&',
        'Or': '|',
        'Xor': '^',
        'Xnor': '^~',
        'Land': '&&',
        'Lor': '||',
        'Sll': '<<',
        'Srl': '>>',
        'Sra': '>>>',
        'Uplus': '+',
        'Uminus': '-',
        'Ulnot': '!',
        'Unot': '~',
        'Uand': '&',
        'Unand': '~&',
        'Uor': '|',
        'Unor': '~|',
        'Uxor': '^',
        'Uxnor': '^~',
    }
    return operator_map.get(op_class_name, op_class_name)



def extract_parameters_from_asts(asts):
    parameters_dict = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                parameters_dict[module_name] = {}
                for item in definition.items:
                    if isinstance(item, Decl):
                        for decl in item.list:
                            if isinstance(decl, Parameter) and not isinstance(decl, Localparam):
                                parameters_dict[module_name][decl.name] = get_rvalue(decl.value)
                    elif isinstance(item, Parameter) and not isinstance(item, Localparam):
                        parameters_dict[module_name][item.name] = get_rvalue(item.value)
    return parameters_dict


def print_parameters(parameters_dict):
    """
    Prints the parameters for each module in CSV format.

    :param parameters_dict: Dictionary with module names as keys and their parameters as values.
    """
    for module_name, params in parameters_dict.items():
        print(f"\nRaw Parameters for module {module_name}:\n")
        if params:
            # Print CSV header
            print("Name,Value")
            for param_name, param_value in params.items():
                print(f"{param_name},{param_value}")
        else:
            print("No parameters found.")


def extract_ioports_from_asts(asts):
    """
    Extracts IO ports' information from the ASTs.
    Returns a dictionary with module names as keys and a list of dictionaries with keys 'Name', 'Width', 'Direction'.

    :param asts: List of ASTs corresponding to the Verilog files.
    :return: Dictionary with module names as keys and their IO ports' information as values.
    """
    ioports_dict = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                ioports_list = []
                for item in definition.items:
                    if isinstance(item, Decl):
                        for decl in item.list:
                            if isinstance(decl, (Input, Output, Inout)):
                                port_direction = decl.__class__.__name__.lower()
                                port_name = decl.name
                                port_width = extract_port_width(decl)
                                width = port_width if port_width else '1'  # If no width specified, assume width is '1'
                                # Append a dictionary
                                ioports_list.append({
                                    'Name': port_name,
                                    'Width': width,
                                    'Direction': port_direction
                                })
                ioports_dict[module_name] = ioports_list
    return ioports_dict


def extract_port_width(decl):
    """
    Extracts the width of a port as a string.

    :param decl: The declaration node of the port.
    :return: A string representing the width of the port.
    """
    if hasattr(decl, 'width') and decl.width is not None:
        msb = get_rvalue(decl.width.msb)
        lsb = get_rvalue(decl.width.lsb)
        # Try to compute the width numerically
        try:
            msb_int = int(msb)
            lsb_int = int(lsb)
            width = abs(msb_int - lsb_int) + 1
            return str(width)
        except ValueError:
            # msb or lsb is not an integer, return as is
            return f"[{msb}:{lsb}]"
    else:
        return "1"  # Default width for 1-bit signals


def print_ioports(ioports_dict):
    """
    Prints the IO ports' information for each module in the same format as extract_ioports_from_specs.

    :param ioports_dict: Dictionary with module names as keys and their IO ports' information as values.
    """
    for module_name, ioports in ioports_dict.items():
        print(f"\nRaw IO Ports for module {module_name}:\n")
        if ioports:
            # Print CSV header
            print("Name,Width,Direction")
            for port in ioports:
                print(f"{port['Name']},{port['Width']},{port['Direction']}")
        else:
            print("No IO ports found.")

def extract_module_hierarchy(asts):
    # Build a mapping from module names to their AST definitions
    module_definitions = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                module_definitions[module_name] = definition

    # Build a mapping from modules to the modules they instantiate
    instantiated_modules = {}
    module_instantiations = {}
    for module_name, definition in module_definitions.items():
        instantiated = []
        instantiations = []
        for item in definition.items:
            if isinstance(item, InstanceList):
                for instance in item.instances:
                    if isinstance(instance, Instance):
                        child_module_name = instance.module
                        instance_name = instance.name
                        instantiated.append(child_module_name)
                        instantiations.append((child_module_name, instance_name))
        instantiated_modules[module_name] = instantiated
        module_instantiations[module_name] = instantiations

    # Identify top-level modules (modules not instantiated by any other module)
    all_modules = set(module_definitions.keys())
    instantiated_modules_set = set()
    for modules in instantiated_modules.values():
        instantiated_modules_set.update(modules)
    top_level_modules = all_modules - instantiated_modules_set

    # Recursively build the module hierarchy starting from top-level modules
    hierarchy_trees = []
    for top_module_name in top_level_modules:
        root_node = ModuleNode(module_name=top_module_name)
        build_hierarchy_recursive(root_node, module_instantiations, module_definitions)
        hierarchy_trees.append(root_node)

    return hierarchy_trees

def build_hierarchy_recursive(current_node, module_instantiations, module_definitions, visited=None):
    if visited is None:
        visited = set()
    module_name = current_node.module_name
    if module_name in visited:
        # Prevent infinite recursion in case of cyclic dependencies
        return
    visited.add(module_name)
    instantiations = module_instantiations.get(module_name, [])
    for child_module_name, instance_name in instantiations:
        child_node = ModuleNode(module_name=child_module_name, instance_name=instance_name)
        current_node.add_child(child_node)
        # Recursively build the hierarchy for the child module
        build_hierarchy_recursive(child_node, module_instantiations, module_definitions, visited.copy())

def print_module_hierarchy(hierarchy_trees):
    for tree in hierarchy_trees:
        tree.print_hierarchy()
        print("")  # Add a blank line between top-level modules


def extract_text_by_chapters(docx_path):
    """
    Extracts text from a Word document and chunks it based on chapter titles,
    including tables, images, and equations, while maintaining the correct order of elements.
    """
    doc = Document(docx_path)
    chapters = []
    current_chapter = {'title': '', 'content': ''}

    image_counter = 1
    equation_counter = 1  # Initialize equation_counter here

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            para = block
            if para.style.name == 'Heading 1':
                if current_chapter['title'] or current_chapter['content']:
                    chapters.append(current_chapter)
                    current_chapter = {'title': '', 'content': ''}  # Start new chapter
                current_chapter['title'] = para.text.strip()
            else:
                if has_equation(para):
                    # Process paragraph with equations
                    equation_counter = process_paragraph_with_equations(para, current_chapter, equation_counter)
                else:
                    # Process runs in the paragraph
                    for run in para.runs:
                        if has_image(run):
                            image_filename = f"image_{image_counter}.png"
                            extract_image(run, image_filename)
                            image_counter += 1
                            current_chapter['content'] += f"[Image: {image_filename}]\n"
                        else:
                            # Add run text
                            current_chapter['content'] += run.text
                    current_chapter['content'] += '\n'
        elif isinstance(block, Table):
            table = block
            # Extract text from the table
            table_text = extract_text_from_table(table)
            # Add table text to the current chapter's content
            current_chapter['content'] += table_text + '\n'
        else:
            # Handle other types of blocks if needed
            pass
    # Add the last chapter
    if current_chapter['title'] or current_chapter['content']:
        chapters.append(current_chapter)
    return chapters


def has_equation(para):
    """
    Checks if the paragraph contains an equation, either as oMath elements or oMathPara elements.
    """
    # Define namespaces
    MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

    # Define fully qualified element names
    oMath_tag = f'{{{MATH_NS}}}oMath'
    oMathPara_tag = f'{{{MATH_NS}}}oMathPara'

    oMath_elements = para._element.findall('.//' + oMath_tag)
    oMathPara_elements = para._element.findall('.//' + oMathPara_tag)

    return bool(oMath_elements or oMathPara_elements)


def process_paragraph_with_equations(para, current_chapter, equation_counter):
    """
    Processes a paragraph containing equations, extracting text and equations
    in order and updating the current chapter's content.
    """
    MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    oMath_tag = f'{{{MATH_NS}}}oMath'
    oMathPara_tag = f'{{{MATH_NS}}}oMathPara'

    content_elements = []

    for child in para._element.iterchildren():
        if child.tag in (oMath_tag, oMathPara_tag):
            # It's an equation
            equation_filename = f"equation_{equation_counter}.xml"
            with open(equation_filename, 'w', encoding='utf-8') as f:
                xml_content = etree.tostring(child, pretty_print=True, encoding='unicode')
                f.write(xml_content)
            content_elements.append(f"[Equation: {equation_filename}]")
            equation_counter += 1
        else:
            # It's text
            text = ''.join(child.itertext())
            content_elements.append(text)

    # Combine the content elements into the current chapter's content
    current_chapter['content'] += ''.join(content_elements) + '\n'

    return equation_counter


def has_image(run):
    """
    Checks if the run contains an image.
    """
    blip_elements = run._element.findall('.//a:blip', {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
    if blip_elements:
        return True
    # Check for embedded images in v:imagedata (used for OLE objects)
    imagedata_elements = run._element.findall('.//v:imagedata', {'v': 'urn:schemas-microsoft-com:vml'})
    if imagedata_elements:
        return True
    return False


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within parent,
    in document order. Each returned value is an instance of either Table or Paragraph.

    :param parent: The parent object (Document or _Cell).
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Unsupported parent type.")

    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)
        else:
            # Handle other types of elements if needed
            pass


def extract_text_from_table(table):
    """
    Extracts text from a table and formats it in Markdown table syntax.

    :param table: The Table object.
    :return: A string representation of the table's content in Markdown format.
    """
    md_table = []

    # Extract the headers (assuming the first row contains headers)
    headers = []
    first_row = table.rows[0]
    for cell in first_row.cells:
        cell_text = ''
        for paragraph in cell.paragraphs:
            cell_text += paragraph.text.strip() + ' '
        headers.append(cell_text.strip())

    # Build the header row
    header_row = '| ' + ' | '.join(headers) + ' |'
    md_table.append(header_row)

    # Build the separator row
    separator_row = '| ' + ' | '.join(['---'] * len(headers)) + ' |'
    md_table.append(separator_row)

    # Process the remaining rows
    for row in table.rows[1:]:
        row_cells = []
        for cell in row.cells:
            cell_text = ''
            for paragraph in cell.paragraphs:
                cell_text += paragraph.text.strip() + ' '
            row_cells.append(cell_text.strip())
        data_row = '| ' + ' | '.join(row_cells) + ' |'
        md_table.append(data_row)

    # Join all rows into a single string
    table_text = '\n'.join(md_table)
    return table_text


def extract_image(run, image_filename):
    """
    Extracts an image from a run and saves it to a file.

    :param run: The run containing the image.
    :param image_filename: The filename to save the image as.
    """
    # Define namespaces
    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'v': 'urn:schemas-microsoft-com:vml',
    }

    # Find all 'blip' elements within the run
    blip_elements = run._element.findall('.//a:blip', nsmap)
    for blip in blip_elements:
        rEmbed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rEmbed:
            image_part = run.part.related_parts[rEmbed]
            with open(image_filename, 'wb') as f:
                f.write(image_part.blob)
            return  # Image extracted

    # Check for 'imagedata' elements (used in OLE objects)
    imagedata_elements = run._element.findall('.//v:imagedata', nsmap)
    for imagedata in imagedata_elements:
        rId = imagedata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        if rId:
            image_part = run.part.related_parts[rId]
            with open(image_filename, 'wb') as f:
                f.write(image_part.blob)
            return  # Image extracted

def normalize_string(s):
    """
    Normalizes a string by replacing curly quotes with straight quotes and removing extra whitespace.
    """
    if s is None:
        return ''
    # Replace curly single quotes with straight quotes
    s = s.replace('‘', "'").replace('’', "'")
    # Replace curly double quotes with straight quotes (if applicable)
    s = s.replace('“', '"').replace('”', '"')
    # Replace any non-breaking spaces with regular spaces (optional)
    s = s.replace('\u00A0', ' ')
    # Strip leading and trailing whitespace
    s = s.strip()
    return s

def get_hierarchy_text(node, level=0):
    """
    Generates a text representation of the module hierarchy starting from the given node.

    :param node: ModuleNode instance representing the root of the hierarchy.
    :param level: Current indentation level.
    :return: A string representing the hierarchy.
    """
    indent = "    " * level
    if node.instance_name:
        line = f"{indent}- Instance: {node.instance_name}, Module: {node.module_name}\n"
    else:
        line = f"{indent}Top-Level Module: {node.module_name}\n"
    for child in node.children:
        line += get_hierarchy_text(child, level + 1)
    return line




def is_clock_signal(signal_name):
    """
    Determines if a signal name corresponds to a clock signal based on common naming conventions.
    """
    clock_keywords = ['clk', 'clock']
    signal_name_lower = signal_name.lower()
    return any(keyword in signal_name_lower for keyword in clock_keywords)


class ModuleClockDomainInfo:
    def __init__(self, module_name):
        self.module_name = module_name
        self.clocks_in_module = []    # List of clocks used in always blocks
        self.has_clock_crossing = False
        self.source_clock = None
        self.destination_clock = None





def get_port_name_by_position(module_definitions, module_name, position):
    """
    Retrieves the port name from the module definition based on its position for positional connections.

    :param module_definitions: Dictionary mapping module names to their AST definitions.
    :param module_name: The name of the module whose port name is to be retrieved.
    :param position: The position index of the port in the module's port list.
    :return: The port name if found, otherwise a placeholder name.
    """
    module_def = module_definitions.get(module_name)
    if module_def and module_def.portlist and module_def.portlist.ports:
        if position < len(module_def.portlist.ports):
            port = module_def.portlist.ports[position]
            return port.name
    return f"port_{position}"


def is_reset_signal(signal_name):
    """
    Determines if a signal name corresponds to a reset signal based on common naming conventions.
    """
    reset_keywords = ['rst', 'reset', 'areset', 'sreset']
    signal_name_lower = signal_name.lower()
    return any(keyword in signal_name_lower for keyword in reset_keywords)




class ModuleResetDomainInfo:
    def __init__(self, module_name):
        self.module_name = module_name
        self.resets_in_module = []    # List of resets used in always blocks
        self.has_reset_crossing = False
        self.source_reset = None
        self.destination_reset = None


def extract_input_ports(module_definitions):
    input_ports_dict = {}
    for module_name, definition in module_definitions.items():
        input_ports = set()

        # Process ports from definition.portlist
        if definition.portlist:
            for port in definition.portlist.ports:
                if isinstance(port, Ioport):
                    decl = port.first
                    if isinstance(decl, Input):
                        port_name = decl.name
                        input_ports.add(port_name)
                elif isinstance(port, Port):
                    port_name = port.name
                    direction = find_port_direction(definition.items, port_name)
                    if direction == 'input':
                        input_ports.add(port_name)

        # Process items in definition.items
        for item in definition.items:
            if isinstance(item, Decl):
                for decl in item.list:
                    if isinstance(decl, Input):
                        port_name = decl.name
                        input_ports.add(port_name)

        if DEBUG:
            print(f"Module: {module_name}, Input Ports: {input_ports}")

        input_ports_dict[module_name] = input_ports
    return input_ports_dict


def find_port_direction(items, port_name):
    """
    Finds the direction of a port by searching through the declarations in the module items.

    :param items: List of module items.
    :param port_name: The name of the port.
    :return: The direction of the port ('input', 'output', 'inout') or None if not found.
    """
    for item in items:
        if isinstance(item, Decl):
            for decl in item.list:
                if decl.name == port_name:
                    if isinstance(decl, Input):
                        return 'input'
                    elif isinstance(decl, Output):
                        return 'output'
                    elif isinstance(decl, Inout):
                        return 'inout'
    return None





def extract_signals_from_expression(expr):
    """
    Recursively extracts signal names from an expression.
    """
    signals = set()
    if isinstance(expr, Identifier):
        signals.add(expr.name)
    elif hasattr(expr, 'children'):
        for child in expr.children():
            signals.update(extract_signals_from_expression(child))
    return signals



class StateMachine:
    def __init__(self, module_name):
        self.module_name = module_name
        self.state_var = None
        self.next_state_var = None
        self.states = []  # List of state names, in order
        self.transitions = []  # List of transitions, each as (current_state, condition, next_state)



def extract_state_machines_from_asts(asts, module_definitions):
    state_machines = {}
    target_module = 'rxuart'  # Specify the module you want to process
    for ast in asts:
        for description in ast.description.definitions:
            if hasattr(description, 'name'):
                module_name = description.name
                if module_name != target_module:
                    continue  # Skip modules that are not the target module
                state_machine = StateMachine(module_name)
                module_items = description.items

                # Extract state variable names and state parameters
                extract_state_variables_and_parameters(module_items, state_machine)

                # Collect known variables
                known_vars = collect_known_vars(state_machine, module_definitions)

                # Find state update and state transition always blocks
                find_state_always_blocks(module_items, state_machine, known_vars)

                if state_machine.state_var and state_machine.transitions:
                    state_machines[module_name] = state_machine
    return state_machines




def extract_state_variables_and_parameters(module_items, state_machine):
    for item in module_items:
        if isinstance(item, Decl):
            for decl in item.list:
                if isinstance(decl, Reg):
                    if 'state' in decl.name:
                        if not state_machine.state_var:
                            state_machine.state_var = decl.name
                            if DEBUG:
                                print(f"Found state variable: {state_machine.state_var}")
                        else:
                            state_machine.next_state_var = decl.name
                            if DEBUG:
                                print(f"Found next state variable: {state_machine.next_state_var}")


def find_state_always_blocks(module_items, state_machine, known_vars):
    for item in module_items:
        if isinstance(item, Always):
            if DEBUG:
                print(f"Always block statement type: {type(item.statement)}")
                print_ast_node(item.statement, "")
            assigned_vars = get_assigned_vars(item.statement)
            if DEBUG:
                print(f"Assigned vars in always block: {assigned_vars}")
            if (state_machine.state_var in assigned_vars or
                state_machine.next_state_var in assigned_vars):
                if DEBUG:
                    print(f"Found state transition always block in module {state_machine.module_name}")
                # Start extraction with collect_states=True
                extract_state_transitions(
                    item.statement,
                    state_machine,
                    collect_states=True,
                    known_vars=known_vars
                )


def get_assigned_vars(statement):
    assigned_vars = set()
    if isinstance(statement, Block):
        for stmt in statement.statements:
            assigned_vars.update(get_assigned_vars(stmt))
    elif isinstance(statement, IfStatement):
        assigned_vars.update(get_assigned_vars(statement.true_statement))
        if statement.false_statement:
            assigned_vars.update(get_assigned_vars(statement.false_statement))
    elif isinstance(statement, CaseStatement):
        for case_item in statement.caselist:
            case_stmt = case_item.statement
            assigned_vars.update(get_assigned_vars(case_stmt))
    elif isinstance(statement, (NonblockingSubstitution, BlockingSubstitution, Assign)):
        # Extract identifiers from the left-hand side recursively
        left_vars = extract_identifiers(statement.left)
        assigned_vars.update(left_vars)
    elif hasattr(statement, 'statement'):
        assigned_vars.update(get_assigned_vars(statement.statement))
    elif isinstance(statement, (list, tuple)):
        for stmt in statement:
            assigned_vars.update(get_assigned_vars(stmt))
    else:
        # Handle other types if necessary
        pass
    return assigned_vars


def extract_identifiers(node):
    identifiers = set()
    if isinstance(node, Identifier):
        identifiers.add(node.name)
    elif hasattr(node, 'children'):
        for child in node.children():
            identifiers.update(extract_identifiers(child))
    return identifiers

def collect_known_vars(state_machine, module_definitions):
    """
    Collects known variable names from the state machine and module definitions.
    """
    known_vars = set()

    # Add state variables
    known_vars.add(state_machine.state_var)
    if state_machine.next_state_var:
        known_vars.add(state_machine.next_state_var)

    # Add input, output, inout, reg, and wire names from the module
    module_name = state_machine.module_name
    module_def = module_definitions.get(module_name)
    if module_def:
        for item in module_def.items:
            if isinstance(item, Decl):
                for decl in item.list:
                    if isinstance(decl, (Input, Output, Inout, Reg, Wire)):
                        known_vars.add(decl.name)

        # Collect identifiers from the module's statements
        for item in module_def.items:
            if isinstance(item, Always):
                collect_identifiers_from_statements(item.statement, known_vars)

    return known_vars

def collect_identifiers_from_statements(statement, known_vars):
    if isinstance(statement, Identifier):
        known_vars.add(statement.name)
    elif hasattr(statement, 'children'):
        for child in statement.children():
            collect_identifiers_from_statements(child, known_vars)
    elif isinstance(statement, list):
        for stmt in statement:
            collect_identifiers_from_statements(stmt, known_vars)
    elif isinstance(statement, (Block, IfStatement, CaseStatement,
                                NonblockingSubstitution, BlockingSubstitution, Assign)):
        # Recursively collect from these statement types
        if hasattr(statement, 'statement'):
            collect_identifiers_from_statements(statement.statement, known_vars)
        if hasattr(statement, 'true_statement'):
            collect_identifiers_from_statements(statement.true_statement, known_vars)
        if hasattr(statement, 'false_statement') and statement.false_statement:
            collect_identifiers_from_statements(statement.false_statement, known_vars)
        if hasattr(statement, 'statements'):
            for stmt in statement.statements:
                collect_identifiers_from_statements(stmt, known_vars)
        if hasattr(statement, 'caselist'):
            for case_item in statement.caselist:
                collect_identifiers_from_statements(case_item.statement, known_vars)
        if hasattr(statement, 'left'):
            collect_identifiers_from_statements(statement.left, known_vars)
        if hasattr(statement, 'right'):
            collect_identifiers_from_statements(statement.right, known_vars)
    # Handle other possible node types if necessary


def get_module_definitions(asts):
    module_definitions = {}
    for ast in asts:
        for definition in ast.description.definitions:
            if hasattr(definition, 'name'):
                module_name = definition.name
                module_definitions[module_name] = definition
    return module_definitions

def extract_state_transitions(statement, state_machine, current_state=None, condition_expr=None, collect_states=False, known_vars=None):
    """
    Extracts state transitions from the given statement and updates the state_machine object.

    :param statement: The AST node representing the statement to process.
    :param state_machine: The StateMachine object to update.
    :param current_state: The current state being processed.
    :param condition_expr: The current condition as a SymPy expression.
    :param collect_states: Boolean indicating whether to collect state names.
    :param known_vars: Set of known variable names for accurate variable/constant identification.
    """
    if known_vars is None:
        known_vars = set()

    if isinstance(statement, CaseStatement):
        comp_value = get_rvalue(statement.comp)
        if DEBUG:
            print(f"Processing CaseStatement with comp: {comp_value}")
        # Check if this is the top-level case statement on the state variable
        if comp_value == state_machine.state_var:
            if DEBUG:
                print("Top-level case statement on state variable")
            collect_states = True
            for case_item in statement.caselist:
                case_conditions = case_item.cond
                if not isinstance(case_conditions, list):
                    case_conditions = [case_conditions]
                for case_cond in case_conditions:
                    if case_cond is None:
                        state_name = 'default'
                        if DEBUG:
                            print("Processing default case")
                    else:
                        state_name = get_rvalue(case_cond)
                    if not isinstance(state_name, str):
                        state_name = str(state_name)
                    if collect_states and state_name != 'default' and state_name not in [state_machine.state_var, state_machine.next_state_var]:
                        if state_name not in state_machine.states:
                            state_machine.states.append(state_name)
                            if DEBUG:
                                print(f"Found state: {state_name}")
                        else:
                            if DEBUG:
                                print(f"State {state_name} already in list.")
                    # Determine the next_state for recursive calls
                    next_state = state_name if state_name != 'default' else current_state
                    # Continue processing with the same condition_expr
                    extract_state_transitions(
                        case_item.statement,
                        state_machine,
                        current_state=next_state,
                        condition_expr=condition_expr,
                        collect_states=False,
                        known_vars=known_vars
                    )
        else:
            if DEBUG:
                print("Nested case statement or case on another variable")
            # In nested case statements, do not update current_state
            for case_item in statement.caselist:
                case_conditions = case_item.cond
                if not isinstance(case_conditions, list):
                    case_conditions = [case_conditions]
                for case_cond in case_conditions:
                    if case_cond is None:
                        # Skip default case in nested case statements
                        if DEBUG:
                            print("Skipping default case in nested case statement")
                        continue  # Skip processing this case
                    # Convert the case condition to SymPy expression
                    case_condition_expr = ast_to_sympy(case_cond)
                    # Combine with existing condition
                    equality_expr = sympy.Eq(ast_to_sympy(statement.comp), case_condition_expr)
                    if condition_expr:
                        combined_condition_expr = SymAnd(condition_expr, equality_expr)
                    else:
                        combined_condition_expr = equality_expr
                    # Process the case item's statement
                    extract_state_transitions(
                        case_item.statement,
                        state_machine,
                        current_state=current_state,
                        condition_expr=combined_condition_expr,
                        collect_states=False,
                        known_vars=known_vars
                    )
    elif isinstance(statement, IfStatement):
        new_condition_expr = ast_to_sympy(statement.cond)
        if condition_expr:
            combined_condition_expr = SymAnd(condition_expr, new_condition_expr)
        else:
            combined_condition_expr = new_condition_expr
        # Process the true branch
        extract_state_transitions(
            statement.true_statement,
            state_machine,
            current_state,
            condition_expr=combined_condition_expr,
            collect_states=False,
            known_vars=known_vars
        )
        # Process the false branch
        if statement.false_statement:
            negated_new_condition_expr = SymNot(new_condition_expr)
            if condition_expr:
                combined_false_condition_expr = SymAnd(condition_expr, negated_new_condition_expr)
            else:
                combined_false_condition_expr = negated_new_condition_expr
            extract_state_transitions(
                statement.false_statement,
                state_machine,
                current_state,
                condition_expr=combined_false_condition_expr,
                collect_states=False,
                known_vars=known_vars
            )
    elif isinstance(statement, (NonblockingSubstitution, BlockingSubstitution, Assign)):
        assigned_var = get_lvalue(statement.left)
        if assigned_var == state_machine.next_state_var or assigned_var == state_machine.state_var:
            next_state_name = get_rvalue(statement.right)
            if not isinstance(next_state_name, str):
                next_state_name = str(next_state_name)
            if next_state_name not in [state_machine.state_var, state_machine.next_state_var]:
                if current_state is not None:
                    # Simplify the condition expression
                    simplified_condition = simplify_condition(condition_expr, known_vars=known_vars) if condition_expr else 'default'
                    state_machine.transitions.append((current_state, simplified_condition, next_state_name))
                    if DEBUG:
                        print(f"Added transition: {current_state} --({simplified_condition})--> {next_state_name}")
                else:
                    if DEBUG:
                        print(f"Skipped transition from None to {next_state_name}")
            else:
                if DEBUG:
                    print(f"Skipped adding transition to variable: {next_state_name}")
    elif isinstance(statement, Block):
        for stmt in statement.statements:
            extract_state_transitions(
                stmt,
                state_machine,
                current_state,
                condition_expr=condition_expr,
                collect_states=False,
                known_vars=known_vars
            )
    elif isinstance(statement, list):
        for stmt in statement:
            extract_state_transitions(
                stmt,
                state_machine,
                current_state,
                condition_expr=condition_expr,
                collect_states=False,
                known_vars=known_vars
            )
    elif hasattr(statement, 'statement'):
        extract_state_transitions(
            statement.statement,
            state_machine,
            current_state,
            condition_expr=condition_expr,
            collect_states=False,
            known_vars=known_vars
        )
    else:
        pass  # Handle other statement types if necessary



def simplify_condition(sympy_expr, known_vars=None):
    """
    Simplifies the SymPy Boolean expression and converts it into a human-readable string.
    """
    simplified_expr = sympy.simplify_logic(sympy_expr)
    return sympy_expr_to_str(simplified_expr, known_vars=known_vars)



def negate_condition(sympy_expr):
    """
    Negates the SymPy Boolean expression and returns a simplified string representation.
    """
    negated_expr = SymNot(sympy_expr)
    simplified_expr = sympy.simplify_logic(negated_expr)
    return sympy_expr_to_str(simplified_expr)


def ast_to_sympy(condition_expr):
    """
    Converts a Verilog AST condition expression into a SymPy Boolean expression.
    """
    if isinstance(condition_expr, Land):
        left = ast_to_sympy(condition_expr.left)
        right = ast_to_sympy(condition_expr.right)
        return SymAnd(left, right)
    elif isinstance(condition_expr, Lor):
        left = ast_to_sympy(condition_expr.left)
        right = ast_to_sympy(condition_expr.right)
        return SymOr(left, right)
    elif isinstance(condition_expr, (Ulnot, Unot)):
        operand = condition_expr.right if hasattr(condition_expr, 'right') else condition_expr.var if hasattr(condition_expr, 'var') else condition_expr.expr
        sym_operand = ast_to_sympy(operand)
        return SymNot(sym_operand)
    elif isinstance(condition_expr, (And, Or)):
        left = ast_to_sympy(condition_expr.left)
        right = ast_to_sympy(condition_expr.right)
        return SymAnd(left, right) if isinstance(condition_expr, And) else SymOr(left, right)
    elif isinstance(condition_expr, Eq):
        left = ast_to_sympy(condition_expr.left)
        right = ast_to_sympy(condition_expr.right)
        return sympy.Eq(left, right)
    elif isinstance(condition_expr, NotEq):
        left = ast_to_sympy(condition_expr.left)
        right = ast_to_sympy(condition_expr.right)
        return SymNot(sympy.Eq(left, right))
    elif isinstance(condition_expr, Identifier):
        return symbols(condition_expr.name)
    elif isinstance(condition_expr, IntConst):
        value = condition_expr.value
        # Preserve the original bit string as a symbol
        return sympy.Symbol(value)
    else:
        # Fallback
        return symbols(get_rvalue(condition_expr))



def sympy_expr_to_str(expr, in_eq=False, known_vars=None):
    """
    Converts a SymPy Boolean expression into a human-readable string.
    Ensures equality expressions are formatted as 'variable is constant'.
    """
    if known_vars is None:
        known_vars = set()

    if expr == S.true:
        return 'True'
    elif expr == S.false:
        return 'False'
    elif isinstance(expr, sympy.Symbol):
        name = expr.name
        # Check if the symbol is a known variable
        if name in known_vars:
            return name if in_eq else f"{name} is HIGH"  # It's a variable
        else:
            return name  # It's a constant
    elif isinstance(expr, sympy.Integer):
        return str(expr)
    elif isinstance(expr, SymNot):
        operand_str = sympy_expr_to_str(expr.args[0], in_eq=in_eq, known_vars=known_vars)
        if operand_str.endswith("is HIGH"):
            return operand_str.replace("is HIGH", "is LOW")
        elif operand_str.endswith("is LOW"):
            return operand_str.replace("is LOW", "is HIGH")
        else:
            return f"{operand_str} is LOW"
    elif isinstance(expr, SymAnd):
        return ' and '.join([sympy_expr_to_str(arg, known_vars=known_vars) for arg in expr.args])
    elif isinstance(expr, SymOr):
        return ' or '.join([sympy_expr_to_str(arg, known_vars=known_vars) for arg in expr.args])
    elif isinstance(expr, sympy.Eq):
        left, right = expr.args

        # Determine if left and right are variables
        left_is_var = is_variable(left, known_vars)
        right_is_var = is_variable(right, known_vars)

        # Swap sides if left is constant and right is variable
        if not left_is_var and right_is_var:
            left, right = right, left

        left_str = sympy_expr_to_str(left, in_eq=True, known_vars=known_vars)
        right_str = sympy_expr_to_str(right, in_eq=True, known_vars=known_vars)
        return f"{left_str} is {right_str}"
    else:
        return str(expr)

def is_variable(sym_expr, known_vars):
    if isinstance(sym_expr, sympy.Symbol):
        return sym_expr.name in known_vars
    else:
        return False





def get_condition_str(condition):
    """
    Converts a Verilog AST condition expression into a human-readable string representation.
    """
    sympy_expr = ast_to_sympy(condition)
    return simplify_condition(sympy_expr)


def get_assigned_states(statement, var_name):
    next_states = []
    if isinstance(statement, (NonblockingSubstitution, BlockingSubstitution)):
        if isinstance(statement.left, Identifier) and statement.left.name == var_name:
            next_state_value = get_rvalue(statement.right)
            next_states.append(next_state_value)
    elif isinstance(statement, Block):
        for stmt in statement.statements:
            next_states.extend(get_assigned_states(stmt, var_name))
    elif isinstance(statement, IfStatement):
        next_states.extend(get_assigned_states(statement.true_statement, var_name))
        if statement.false_statement:
            next_states.extend(get_assigned_states(statement.false_statement, var_name))
    elif isinstance(statement, CaseStatement):
        for case_item in statement.caselist:
            next_states.extend(get_assigned_states(case_item.statement, var_name))
    return next_states


# cywong
def compare_state_machines(rtl_state_machine, specs_chapter_content, module_name):
    """
    Compares the state machine extracted from RTL with the specifications.
    Uses GPT to perform the comparison and outputs the result.
    """
    # Prepare the RTL state machine description
    rtl_state_machine_description = generate_rtl_state_machine_description(rtl_state_machine)

    # Prepare the prompt
    prompt = (
        "You are an assistant that compares state machines extracted from RTL code with those described in specifications. "
        "Your goal is to identify any mismatches between them, focusing on state names and signal levels (HIGH or LOW) in transition conditions.\n\n"
        "**Important Notes:**\n\n"
        "- **State Names Must Match Exactly:**\n"
        "  - Any differences in state names between the RTL code and the specifications should be considered a mismatch.\n"
        "  - For example, `RXU_IDLE` and `RXU_IDLE1` are different state names and should be reported as a mismatch.\n\n"
        "- **Signal Levels in Transition Conditions Must Match Exactly:**\n"
        "  - Any differences in signal levels (e.g., HIGH vs. LOW) in transition conditions between the RTL code and the specifications should be considered a mismatch.\n"
        "  - Pay close attention to whether transitions occur on signals being HIGH or LOW.\n\n"
        "- **Logical Equivalence of Conditions:**\n"
        "  - **Do not consider differences in how the logical conditions are structured, including nesting, as a mismatch if the logical conditions are equivalent.**\n"
        "  - **Only report mismatches if the logical conditions are not equivalent or if the signal levels differ.**\n"
        "  - When comparing complex logical expressions, perform step-by-step simplification using logical identities and laws (e.g., De Morgan's laws).\n\n"
        "- **Important Note on Condition Simplification:**\n"
        "  - When comparing transition conditions, **always simplify the RTL and specification conditions step-by-step** to their basic signal levels.\n"
        "  - **Only consider the final simplified conditions** when checking for mismatches.\n"
        "  - **Do not report mismatches based on differences in expression structure, grouping, or use of parentheses**, as long as the signal levels are equivalent.\n\n"
        "- **Example:**\n"
        "  - **Example 1:**\n"
        "    - **RTL Condition:** `o_break is LOW and (ck_uart is HIGH or half_baud_time is LOW)`\n"
        "    - **Specification Condition:** `o_break is LOW and (ck_uart is HIGH or half_baud_time is LOW)`\n"
        "    - **Explanation:** Both conditions are logically equivalent and should be considered matching.\n"
        "  - **Example 2:**\n"
        "    - **RTL Condition:** `o_break is LOW and zero_baud_counter is HIGH and ck_uart is LOW`\n"
        "    - **Specification Condition:** `o_break is LOW and zero_baud_counter is HIGH and ck_uart is LOW`\n"
        "    - **Explanation:** Both conditions are logically equivalent and should be considered matching.\n\n"
        "- **Signal Interpretation:**\n"
        "  - Signals can be active HIGH or active LOW, as specified.\n\n"
        "  - **Active Levels of Signals:**\n"
        "    - `line_synch` is an active HIGH signal indicating line synchronization.\n"
        "      - When `line_synch` is **HIGH**, synchronization is achieved.\n"
        "      - When `line_synch` is **LOW**, synchronization is not achieved.\n"
        "    - `o_break` is an active HIGH signal indicating a break condition.\n"
        "      - When `o_break` is **HIGH**, a break condition is active.\n"
        "      - When `o_break` is **LOW**, the break condition has ended.\n"
        "    - `ck_uart` is an active HIGH signal representing the UART line state.\n"
        "      - When `ck_uart` is **HIGH**, the line is idle or in a stop bit.\n"
        "      - When `ck_uart` is **LOW**, the line is in a start bit or transmitting data.\n"
        "    - `zero_baud_counter`: Signal that is **HIGH** when the baud counter reaches zero.\n"
        "    - `dblstop`: Signal that is **HIGH** when double stop bits are enabled.\n"
        "    - `use_parity`: Signal that is **HIGH** when parity is used.\n"
        "    - `data_bits`: A 2-bit signal indicating the number of data bits:\n"
        "      - `2'b00`: 5 data bits\n"
        "      - `2'b01`: 6 data bits\n"
        "      - `2'b10`: 7 data bits\n"
        "      - `2'b11`: 8 data bits\n"
        "    - `half_baud_time`: Signal that is **HIGH** at half the baud rate interval.\n\n"
        "**Your Task:**\n\n"
        "- Compare the state names in the RTL code and the specifications. Any difference in state names should be considered a mismatch.\n"
        "- Compare the transition conditions, paying close attention to signal levels (e.g., HIGH vs. LOW).\n"
        "- **Do not consider differences in how the logical conditions are structured, including nesting, as a mismatch if they are logically equivalent.**\n"
        "- **Only report mismatches if state names differ or if signal levels in the transition conditions differ.**\n"
        "- Identify any mismatches in state names, signal levels, or transition conditions.\n"
        "- Report any discrepancies, including differences in state names, signal levels, and transition conditions.\n"
        "- **State names and explicit signal levels must match exactly.**\n\n"
        "- **Do not consider differences in the structure of conditions (nested vs. flat) as a mismatch if the logical conditions are equivalent, but state names and signal levels must match exactly.**\n\n"
        "**Note:**\n\n"
        "- If they match in all aspects, output:\n\n"
        "  [INFO ] MATCHED\n\n"
        "- If they do not match, output:\n\n"
        "  [ERROR] MISMATCHED. Provide the precise reason for the mismatch, including specific differences in state names, signal levels in transitions, or conditions.\n\n"
        "Only output the specified message, without any additional text.\n\n"
        "**State Machines for Module `{}`:**\n\n"
        "**RTL State Machine:**\n\n"
        "{}\n\n"
        "**Specifications State Machine:**\n\n"
        "{}"
    ).format(module_name, rtl_state_machine_description.strip(), specs_chapter_content.strip())

    if DEBUG:
        print("GPT Prompt:\n", prompt)

    # Call GPT to perform the comparison
    try:
        completion = client.chat.completions.create(
            model='gpt-4o',
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an assistant that compares state machines from RTL code and specifications. "
                        "You only output the comparison result precisely as specified."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=2000,
            n=1,
            stop=None,
        )

        # Get the assistant's reply
        assistant_reply = completion.choices[0].message.content.strip()

        # Print the assistant's reply
        print(assistant_reply)

        # Return the result for final summary
        if "[INFO ]" in assistant_reply:
            result = assistant_reply.replace("[INFO ]", "").strip()
            return ("State Machine Check", result)
        elif "[ERROR]" in assistant_reply:
            result = assistant_reply.replace("[ERROR]", "").strip()
            return ("State Machine Check", result)
        else:
            return ("State Machine Check", assistant_reply)

    except Exception as e:
        print(f"[ERROR] Error comparing state machines from specifications and RTL: {e}")
        return ("State Machine Check", "Error during comparison.")


def generate_rtl_state_machine_description(state_machine):
    """
    Generates a textual description of the RTL state machine for inclusion in the GPT prompt
    """
    description = f"State Variable: {state_machine.state_var}\n"
    description += f"Next State Variable: {state_machine.next_state_var}\n\n"
    description += "States:\n"
    for state_name in state_machine.states:
        description += f"  {state_name}\n"
    description += "\nTransitions:\n"
    for current_state, condition_expr, next_state in state_machine.transitions:
        # The condition_expr is already a simplified, readable string
        description += f"  From {current_state} -> {next_state} on {condition_expr}\n"
    return description


def print_single_state_machine_info(state_machine):
    if state_machine is None:
        print("No state machine information available.")
        return

    print(f"\nState Machine in Module: {state_machine.module_name}")
    print(f"State Variable: {state_machine.state_var}")
    print(f"Next State Variable: {state_machine.next_state_var}")
    print("\nStates:")
    for state_name in state_machine.states:
        print(f"  {state_name}")
    print("\nTransitions:")
    for current_state, condition, next_state in state_machine.transitions:
        print(f"  From {current_state} -> {next_state} on {condition}")
    print("\n")



def main():
    run_state_machine_check = True

    # Path to your specifications Word document
    spec_docx_path = 'specs/specs.docx'

    # # List of modules' name
    module_names = ['ufifo', 'rxuart', 'txuart', 'wbuart']

    # Get the list of Verilog files
    file_list = get_file_list("rtl")

    # Generate ASTs for the files
    asts = generate_asts(file_list)

    # Prints the AST contents for each module.
    if DEBUG:
        print_ast_contents(asts)

    # Collect module definitions
    module_definitions = get_module_definitions(asts)

    # Extract chapters from the Word document
    chapters = extract_text_by_chapters(spec_docx_path)
    final_results = []  # Store the results as tuples (check_name, result_content)

    # Find the chapter with title 'State Machine'
    state_machine_chapter = next((chapter for chapter in chapters if 'state machine' in chapter['title'].lower()), None)

    if run_state_machine_check:
        if state_machine_chapter:
            # Add header for State Machine Check
            print("\n===== State Machine Check =====\n")

            # Extract state machine from RTL, passing module_definitions
            rtl_state_machines = extract_state_machines_from_asts(asts, module_definitions)

            # # Assuming we're focusing on the 'rxuart' module
            module_name = 'rxuart'
            rtl_state_machine = rtl_state_machines.get(module_name, None)

            if rtl_state_machine:
                print_single_state_machine_info(rtl_state_machine)

                # Get the specifications chapter content
                specs_chapter_content = state_machine_chapter['content']

                # Compare State Machines
                final_results.append(compare_state_machines(rtl_state_machine, specs_chapter_content, module_name))
            else:
                print(f"[ERROR] State machine information not found in RTL for module {module_name}.")
        else:
            print("[ERROR] Chapter 'State Machine' not found in the specifications.")


    # Calculate the maximum length of check names for alignment
    if final_results:
        max_check_name_length = max(len(check[0]) for check in final_results)

        # Print all results with aligned check names
        print("\n===== Final Results =====")
        for check_name, result_content in final_results:
            print(f"{check_name.ljust(max_check_name_length)} : {result_content}")


if __name__ == "__main__":
    main()

